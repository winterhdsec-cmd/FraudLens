#!/usr/bin/env python3
"""Build CPEC2026 teaching-case paper as DOCX matching the reduced LaTeX content."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = r'E:\FraudLens\paper\论文产出\CPEC2026教学案例稿'


def set_default_font(doc, font_name='Times New Roman', cjk_name='SimSun'):
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(12)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), cjk_name)


def add_centered_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    if level == 0:
        run.font.size = Pt(18)
        run.bold = True
    elif level == 1:
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    return p


def add_left_heading(doc, text, level=2):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(14 if level == 2 else 12)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(6)
    return p


def add_para(doc, text, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, space_before=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_label_para(doc, label, content, bold_label=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    r1.bold = bold_label
    r2 = p.add_run(content)
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p


def add_figure(doc, png_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    crun.font.name = 'Times New Roman'
    crun._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    crun.font.size = Pt(10.5)
    crun.bold = True


def add_table(doc, rows, caption, col_widths=None, header_bold=True):
    # caption above table
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    crun.font.name = 'Times New Roman'
    crun._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    crun.font.size = Pt(10.5)
    crun.bold = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    run.font.size = Pt(10.5)
                    if i == 0 and header_bold:
                        run.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    set_default_font(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('面向反诈团伙研判的 AI 赋能实验教学案例设计')
    tr.font.name = 'Times New Roman'
    tr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    tr.font.size = Pt(18)
    tr.bold = True
    t.paragraph_format.space_after = Pt(6)

    # Authors
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = a.add_run('韩冬，吴燕波*，徐伟')
    ar.font.name = 'Times New Roman'
    ar._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    ar.font.size = Pt(12)

    # Affiliation
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affr = aff.add_run('（湖北警官学院 信息技术系，武汉 430034，中国）')
    affr.font.name = 'Times New Roman'
    affr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    affr.font.size = Pt(10.5)

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    metar = meta.add_run('收稿日期：2026-XX-XX；基金项目：湖北省大学生创新创业训练计划项目（编号：S202611332001，徐伟、吴燕波共同指导）；院级科研项目（编号：待补）。\n'
                         '作者简介：韩冬（2005—），男，湖北，本科在读，主要研究方向为网络安全与人工智能反诈，E-mail：winterhdsec@163.com（电话：18202799140）。\n'
                         '徐伟（1978—），男，湖北，博士在读，副教授，研究方向为网络空间安全（网络入侵检测、物联网安全），E-mail：8073@hbpa.edu.cn。\n'
                         '通信作者：吴燕波（19xx—），女，湖北，硕士，副教授，研究方向为公安院校计算机与网络安全实践教学、网络安全与执法人才培养、警用智能装备应用，E-mail：240325743@qq.com（电话：13339993182）。')
    metar.font.name = 'Times New Roman'
    metar._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    metar.font.size = Pt(9)

    # Chinese abstract
    add_centered_heading(doc, '摘  要', level=1)
    add_para(doc,
        '针对公安院校反诈实战教学中学生难以动手操作真实研判系统、教师缺乏可直接落地的现成案例等问题，本文以一个可运行的反诈研判原型系统 FraudLens 为载体，设计了一套学生可一键复现的 AI 赋能实验教学案例。'
        '案例将系统的“串并案—扩线—反思”两级研判闭环转化为实训内容，以“研—训—评”模式组织四个递进实验环节，'
        '引导学生在 Jupyter 环境中开展人机协同研判与算法边界认知训练，配套含权重的评价量表与思政映射。'
        '案例以研判决策而非模型分析为核心，旨在培养学生在 AI 辅助下的领域研判决策、人机协同与工程可信伦理素养，'
        '为公安实战人才培养提供可复用的领域级实训素材。', indent=False)
    add_label_para(doc, '关键词：', '反诈教学；图神经网络；实验教学案例；研判决策；人机协同')
    add_label_para(doc, '中图分类号：', 'TP393.08；TP18；G642    文献标志码：A')

    # English header
    et = doc.add_paragraph()
    et.alignment = WD_ALIGN_PARAGRAPH.CENTER
    etr = et.add_run('Design and Exploration of an AI-Enabled Practical Teaching Case for Anti-Fraud')
    etr.font.name = 'Times New Roman'
    etr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    etr.font.size = Pt(16)
    etr.bold = True

    ea = doc.add_paragraph()
    ea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ear = ea.add_run('HAN Dong, WU Yanbo*, XU Wei\nHubei University of Police, Department of Information Technology, Wuhan 430034, China')
    ear.font.name = 'Times New Roman'
    ear._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    ear.font.size = Pt(12)

    add_left_heading(doc, 'Abstract:', level=3)
    add_para(doc,
        'To address the gap that students in anti-fraud practical teaching at police colleges can hardly operate a real investigation system hands-on while instructors lack ready-to-use, reproducible case material, '
        'this paper presents a reproducible, AI-enabled practical teaching case built on a working anti-fraud investigation prototype, FraudLens. '
        'The case reformulates the system\'s two-level investigation loop (case-level merging and account-level expansion) and its reflection orchestration into hands-on laboratory content, '
        'and organizes a four-lab progressive chain—case analysis, AI-assisted case linking, freeze-order ethical decision-making, and capability-boundary reflection—'
        'under a research–training–evaluation model. Guided by one-click reproducible baselines, students conduct human–AI collaborative investigation and algorithm-boundary reflection in a Jupyter environment, '
        'with weighted rubrics and ideological–ethical mapping throughout. Centered on investigative decision-making rather than model analysis, the case '
        'cultivates students\' AI-assisted decision-making, human–AI collaboration, engineering trustworthiness, ethical responsibility, and boundary awareness, '
        'thereby providing a deployable domain-level training resource for police talent development.',
        indent=False)
    add_label_para(doc, 'Key words: ', 'anti-fraud; graph neural network; practical teaching case; investigative decision-making; human–AI collaboration')

    # Section 1
    add_centered_heading(doc, '1  引言', level=1)
    add_para(doc,
        '新工科建设对计算机类人才提出了 AI 赋能实践能力与工程素养的要求，将真实领域问题转化为可复现的实践教学案例，'
        '是提升学生实战能力的重要途径[1,2]。然而，现有计算机实践教学案例多聚焦通用编程与单门课程建设[2]，'
        '面向领域级实战研判（如涉网犯罪智能研判）的可复现案例仍属稀缺[3]；'
        '尤其对公安院校而言，反诈实战教学长期面临“学生只能听案例、看演示，却难以亲手操作一套真实研判系统，教师也缺乏可直接用于课堂、又能让学生复现的现成案例”的困境。')
    add_para(doc,
        '需要说明的是，本文是一篇面向反诈实战教学的实验教学案例设计，而非算法研究论文；其落脚点在于如何将可运行的反诈研判原型系统 FraudLens 转化为学生可动手、教师可直接落地的实训载体，而非展开算法本身的实现细节。')
    add_para(doc,
        '按照工程教育对复杂工程问题的通行界定（须深入工程原理分析、多因素冲突、需建立抽象模型、非常规方法可解、利益不一致、含多个关联子问题）[3,5]，'
        '反诈团伙研判可逐条对应：资金拓扑与话术文本构成多源异构数据，需建立图抽象模型；团伙反侦查行为引入多因素冲突；'
        '冻卡决策涉及利益与规范权衡；研判流程包含数据建模、表示学习、系统编排与可信决策等多个关联子问题。'
        '因此，以反诈团伙研判为载体的实训，能让学生在真实问题中综合运用图建模、深度表示学习、系统编排与工程可信设计；'
        '其教学落脚点并非让学生独立“解决复杂工程问题”，而是培养其在 AI 辅助下做出领域研判决策、并对算法能力边界保持审慎认知的能力。')
    add_para(doc,
        '基于上述认识，本文以可运行的反诈团伙研判原型系统 FraudLens 为载体，将“案件级串并案 + 账户级扩线”两级研判闭环与反思编排整合为实训内容，构建可复现的“研—训—评”一体化教学案例。本文主要工作如下：')

    works = [
        '可落地的 AI 辅助研判决策实训载体：FraudLens 两级研判系统（“案件级串并案 + 账户级扩线”闭环，由 LangGraph StateGraph[6] 编排“规划 → 预处理 → 分析 → 聚类 → 反思”工作流），技术栈覆盖图建模、异构注意力、系统编排与工程可信，可作为学生“研”的对象逐层复现。',
        '“研—训—评”一体化的实训模式：按“案情分析 → 工具辅助串并案 → 冻卡决策与伦理权衡 → 边界认知与反思”设计 4 环节递进实验链（第 5 节），以研判决策而非模型分析为核心活动，配套含权重的评价量表与思政/伦理映射。',
        '可复现实验与科学边界认知训练：能力基线（HAN 困难场景 F1=0.9154、反思闭环增益 +0.5125、账户级扩线将团伙 F1 由盲扫约 0.002 提升至约 0.71）全部可一键复现；同时诚实量化反诈 GNN 的增量边界（合成 ≠ 真实、盲扫全败、不构成真实警务验证），引导学生建立对模型能力的审慎认知。'
    ]
    for w in works:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r = p.add_run(w)
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r.font.size = Pt(12)
    add_para(doc,
        '需要说明的是，文中能力基线数据（如 HAN 困难场景 F1、反思闭环增益、扩线召回）仅作为实训可复现对象的佐证，不代表真实警务场景的已验证成效；教学成效的课堂实证评估属下一阶段工作。')

    # Section 2
    add_centered_heading(doc, '2  相关工作', level=1)
    add_left_heading(doc, '2.1  反诈团伙检测技术', level=2)
    add_para(doc,
        '反诈检测经历了从规则到图神经网络的演进：基于规则的串并案依赖办案经验、难以规模化；浅层聚类无法建模账户间复杂关联；'
        '图神经网络（GNN）[7] 虽能利用图结构，欺诈场景亦已考虑伪装节点与拓扑[4]，但同构 GNN 难以同时建模案件、账户、违法者等多类实体及其关系——'
        '这正是本文采用异构注意力网络（HAN）[8] 的动机。从教学角度看，这条技术演进线本身即“数据建模 → 表示学习 → 工程决策”的完整教材。')
    add_left_heading(doc, '2.2  人工智能赋能实验教学', level=2)
    add_para(doc,
        '以大语言模型为代表的人工智能正加速融入计算机实践教学：李清勇等[1] 探讨了通用大模型在实践教学中的角色定位，'
        '张金等[2] 构建了基于通用大模型的系统创新实验。然而，现有案例多聚焦通用编程与单门课程，'
        '面向领域级实战研判的可复现案例仍属稀缺[3]。本文据此将 FraudLens 设计为可复现实验案例，支撑学生从数据建模到科学边界认知的递进训练。')

    # Section 3
    add_centered_heading(doc, '3  实训载体：FraudLens 系统概览', level=1)
    add_para(doc,
        'FraudLens 是一套可运行的反诈团伙研判原型系统（图 1），覆盖数据建模、表示学习、系统编排与工程可信的完整技术链，'
        '恰好对应实训所需的知识与能力结构，是理想的“研”对象。系统逻辑上构成两级研判闭环：案件级串并案（一级）以 HAN 在案件异构图上聚合同伙案件，'
        '账户级扩线（二级）在给定线索锚点后于账户交易子图上还原资金链路；一级输出的关键账户可直接作为二级扩线的锚点，契合真实反诈工作流。'
        '对实训而言，这套闭环覆盖了从数据到决策的完整链路，学生可沿两级闭环逐段复现。')
    add_figure(doc, os.path.join(BASE, 'fig_arch.png'), '图 1  FraudLens 系统总体架构')
    add_para(doc,
        '为降低上手门槛，实训以研判决策为核心，不要求学生理解图神经网络内部原理；全部模块均提供预计算结果供学生交互使用。系统的技术构成可概括为四点：'
        '（1）异构双通道融合：以 HAN[8] 融合“结构通道（金额加权交易图）”与“文本通道（本地 BGE 语义嵌入）”[9]，支撑案件级团伙聚类；'
        '（2）多智能体反思闭环：由 LangGraph[6] 驱动“规划 → 预处理 → 分析 → 聚类 → 反思”工作流，反思节点在聚类质量不达标时经条件边回连触发真实重算[10]；'
        '（3）经验加权置信度门控：以团伙规模、涉案金额、关联账户数、资金回流四因子加权给出冻卡建议置信度；'
        '（4）工程可信：实现 RBAC 鉴权、双表审计、多级降级链与“数据不出域”，契合公安院校数据合规要求。上述技术细节并非本文讨论重点，本文仅将其作为实训“研”的对象。')
    add_para(doc,
        '为直观呈现学生如何操作 FraudLens 并与之协同研判，图 2 给出学生实验流程与系统响应机制：学生以四环节实验链为单位向系统输入案情与参数，'
        '系统经预处理、语义嵌入与图构建形成研判对象，由 LangGraph 多智能体状态图完成“规划—预处理—分析—聚类—反思”，输出候选结果供学生进行人机协同决策；'
        '全程审计日志与加权评价量表支撑“研—训—评”闭环。')
    add_figure(doc, os.path.join(BASE, 'fig_workflow.png'), '图 2  学生实验流程与系统响应机制（示意图）')

    # Section 4
    add_centered_heading(doc, '4  实训可复现性与能力基线', level=1)
    add_para(doc,
        '实训以可复现为设计前提：合成案情数据集随系统开放，全部实验提供一键复现脚本；HAN 采用无标签预训练，不依赖人工团伙标注；'
        '外部验证所用 Elliptic[11]、AMLSim[12] 均为公开数据，学生可在本地完整复现。能力基线中的合成数据不等于真实警务数据，这一边界本身即是实训 Lab4 的教学内容。')
    add_left_heading(doc, '4.1  能力基线与消融', level=2)
    add_para(doc,
        '表 1 报告合成数据（10 随机种子均值）下的能力基线。HAN 在困难（Hard）场景 F1=0.9154，显著优于同构 GraphSAGE（Hard F1=0.3043，'
        '因无法建模多类型实体关系发生塌缩），这一对比直接成为实训 Lab2“为何需要 AI 辅助”的教学素材。消融显示，反思闭环仅作用于未收敛的失败场景：'
        '其缺失使失败场景 F1 由 0.8353 跌至 0.3228（增益 +0.5125），说明反思闭环对该类场景的关键作用；关闭 GNN 退化为 Louvain 社区发现，'
        'Hard F1 由 0.9154 降至 0.8616，说明异构图建模对困难场景的必要贡献。')
    add_table(doc,
        [
            ['方法/配置', 'Clean F1', 'Hard F1', '备注'],
            ['HAN（本文）', '1.0000±0.0000', '0.9154±0.1214', '异构注意力融合'],
            ['GraphSAGE', '0.4061±0.2127', '0.3043±0.0000', 'Hard 塌缩（同构局限）'],
            ['Louvain（降级）', '0.8884±0.0911', '0.8616±0.0874', '关闭 GNN 的社区发现'],
            ['w/o 反思闭环', '—', '—', '失败场景 F1 由 0.8353 跌至 0.3228'],
        ],
        '表 1  能力基线与消融（合成数据，10 seeds mean±std）',
        col_widths=[1.4, 1.1, 1.1, 2.4])
    add_para(doc,
        '注：Clean 场景无跨团伙干扰（cross=0.0），仅验证模型基本学习能力，不代表真实场景预期。反思闭环仅作用于失败场景的自动降级重算，'
        '不改变已收敛场景均值，故 w/o 反思的 Hard F1 以“—”表示。', indent=False, space_before=18)

    add_left_heading(doc, '4.2  账户级扩线外部验证（AMLSim）', level=2)
    add_para(doc,
        '为验证账户级扩线（二级闭环）在真实规模公开图上的能力边界，本节在 AMLSim[12] 反洗钱基准（43,614 账户 / 1,305 个洗钱环，含账户级环真值）上报告无监督扩线结果：'
        '给定嫌疑账户锚点，于账户交易图上做 k=1 跳扩线，子图规模约 2,748 节点（环占比约 0.55），全部方法零标签、纯无监督。表 2 显示，'
        '盲扫（全网无锚点）几乎失效（F1≈0.002），印证“盲扫 vs 扩线”瓶颈；扩线设定下，k-core 核心子图（kmin=2）以无监督方式将团伙识别 F1 提升至约 0.71（召回 0.91），'
        '短有向环检测将资金环识别 F1 提升至约 0.62。该结果不宣称真实警务验证——AMLSim 为公开合成基准；其适用边界明确：仅适用于含金额/时序的账户交易图，'
        '对纯拓扑交易图（Elliptic，无金额时序）不迁移。上述“无监督有上限、标注能破但不稳”的诚实结论，正是实训 Lab4 的数据素材。')
    add_table(doc,
        [
            ['任务', '方法', 'Precision', 'Recall', 'F1'],
            ['团伙识别', '盲扫（无锚点）', '—', '—', '≈0.002（失效）'],
            ['团伙识别', 'k-core 核心子图（kmin=2）', '0.58', '0.91', '0.71'],
            ['资金环识别', '短有向环检测（L≤8）', '0.63', '0.61', '0.62'],
        ],
        '表 2  AMLSim 账户级扩线无监督基线（锚点 k=1 跳，零标签）',
        col_widths=[1.2, 2.4, 1.1, 0.9, 1.0])
    add_para(doc,
        '注：AMLSim 为 IBM 公开合成反洗钱基准，含账户级环真值；全部方法零标签纯无监督，不构成真实警务数据验证。'
        'k-core 召回 0.91 但候选占子图约 85%，故 precision 受限——这正是无监督上限的结构性原因（详见 §5.3.4 Lab4）。', indent=False, space_before=18)

    # Section 5
    add_centered_heading(doc, '5  实训教学设计', level=1)
    add_para(doc,
        'FraudLens 所覆盖的技术栈——异构图谱建模、双通道语义融合、多智能体反思编排、可解释门控决策——恰好对应“新工科”对 AI 赋能实践能力与工程素养的要求。'
        '本节给出将其转化为可复现实训教学案例的设计方案，供"人工智能安全""图数据挖掘"及相关专业课程参考。'
        '以下为基于系统真实能力的教学案例设计，课堂实施与成效以实际开课情况为准。')
    add_left_heading(doc, '5.1  学情分析与课程衔接', level=2)
    add_para(doc,
        '本案例面向公安院校信息技术、网络安全与侦查相关专业的本科生（建议大二下至大三上）实施。学生已修 Python 程序设计、数据结构与数据库等课程，'
        '具备基本图论概念与编程调试能力，但一般无图神经网络与多智能体先验[2]；图神经网络与多智能体不作为先修要求，仅以黑箱工具形式参与，'
        '确保实训不依赖高阶先验。课程对接《人工智能安全》《图数据挖掘》等。据此，实训以研判决策为核心活动：'
        '图神经网络作为学生交互的黑箱工具而非需要理解的教学内容，学生只需理解“系统给出什么结果、可信到什么程度、何时该信、何时不该信”，'
        '而非掌握模型内部机理。全部实验以 Jupyter Notebook 为载体，系统研判结果预烘焙供学生对比分析。')

    add_left_heading(doc, '5.2  教学理念与“研—训—评”一体化模式', level=2)
    add_para(doc,
        '本节以“研—训—评”一体化模式组织实训（图 3）：以真实 FraudLens 系统为“研”的对象，以四环节递进实验链为“训”的载体，'
        '以过程+产物双轨评价及思政/伦理素养观测为“评”的手段，并以可复现实验、诚实边界认知与课程思政作为全程支撑。'
        '该模式将技术内容组织为“研究—训练—评价”的教学闭环，避免技术细节淹没育人主线。')
    add_para(doc,
        '需要说明：上述“研—训—评”构成的是教学设计层面的闭环（目标—活动—评价—反思），其结构已完整闭合；'
        '而实证成效层面的闭环——即学生实际走完四环节、测出能力提升并反馈改进——将在 2026 年秋季学期课程实施后闭合，'
        '当前稿以设计论证与可复现能力基线为支撑，教学成效以实际开课情况为准。')
    add_figure(doc, os.path.join(BASE, 'fig_mode.png'), '图 3  “研—训—评”一体化反诈实训模式框架')

    add_left_heading(doc, '5.3  四环节递进实验链', level=2)
    add_para(doc,
        '实训以四环节递进实验链为核心抓手（图 4）：四个 Lab 从案情分析到边界反思逐级递进，每个 Lab 锚定 FraudLens 的真实模块，以研判决策为核心活动。'
        '系统研判结果预烘焙为 JSON/NumPy 数组，学生打开 notebook 执行 Run All 即可获得案情数据与系统分析结果，教学时间用于决策、对比与反思。'
        '表 3 给出各 Lab 的课时分配、核心任务、对应能力目标与训练产出。')
    add_figure(doc, os.path.join(BASE, 'fig_stages.png'), '图 4  四环节递进实验链及其与系统模块对应')
    add_table(doc,
        [
            ['Lab', '时长', '核心任务', '对应能力/系统模块', '训练产出'],
            ['1', '4 课时', '分析案情，认识研判流程', '案情分析能力·§3 系统概览', '案情分析报告'],
            ['2', '4 课时', '人工串并案 vs 系统结果', '人机协同·双通道 HAN+反思闭环', '对比报告'],
            ['3', '4 课时', '门控调参，冻卡决策角色扮演', '决策与伦理·置信度门控', '伦理分析报告'],
            ['4', '4 课时', '系统失败场景与诚实反思', '科学边界认知·外部验证', '边界反思总结'],
        ],
        '表 3  四环节递进实验链——课时分配、核心任务与训练产出',
        col_widths=[0.5, 0.7, 1.7, 1.9, 1.0])
    add_para(doc, '注：课前预习 2 课时（pip install + 数据集下载，不计入课内学时）。课堂导入与总结内嵌于各 Lab，课内总计 16 学时。', indent=False, space_before=18)

    # Lab descriptions
    add_left_heading(doc, '5.3.1  Lab1：案情分析与研判流程', level=3)
    add_para(doc,
        '教学目标：理解反诈案件的数据结构与研判流程。学生获得一份合成案情（含受害者陈述、账户流水、通话记录、涉案金额），以办案人员视角分析：涉及哪些账户？钱怎么转的？'
        '哪个环节是关键节点？然后系统展示其如何将同一案件建模为异构图可视化（案件—账户—手机号—类型—城市）。收尾讨论：“如果只用传统逐笔查账方式，你会漏掉什么？”')
    add_left_heading(doc, '5.3.2  Lab2：工具辅助串并案', level=3)
    add_para(doc,
        '教学目标：体验 AI 辅助串并案的价值与局限，建立“人机协同”意识。学生先人工判断“哪几个案子可能是同一伙人干的”，再运行系统看聚类结果。关键教学时刻：'
        '对比人找到的和系统找到的——系统发现了人漏掉的多跳关联，但人也可能发现系统漏掉的（基于办案经验而非数据特征的关联）。核心学到：AI 能看到人看不到的结构关联，但也有盲区——AI 是辅助工具而非替代。')
    add_left_heading(doc, '5.3.3  Lab3：冻卡决策与伦理权衡', level=3)
    add_para(doc,
        '教学目标：在真实决策压力下理解技术—伦理权衡。学生调整置信度门控阈值，观察阈值变化如何影响冻卡建议数量。角色扮演：“如果阈值设高了，3 个无辜者的账户被冻结；'
        '如果设低了，团伙今晚转移资金。”学生记录决策理由并讨论：“算法建议冻，但谁来承担后果？”核心：冻卡不是技术问题，是伦理问题。')
    add_left_heading(doc, '5.3.4  Lab4：边界认知与诚实反思', level=3)
    add_para(doc,
        '教学目标：建立对 AI 研判工具能力边界的认知。学生观察同构 GNN 基线（GraphSAGE）在困难场景塌缩（F1 由约 0.41 跌至 0.30），而异构 HAN 保持 0.9+；'
        '再读 AMLSim 盲扫 F1≈0.002、Elliptic HAN F1≈0.016——真实数据上所有方法都接近失败。引入“扩线”概念：给定锚点 → 只分析 k 跳邻居 → 无监督拓扑方法将团伙识别 F1 由盲扫 0.002 提升至约 0.71、'
        '资金环识别约 0.62——无需任何标签，纯无监督即可显著恢复团伙与资金链；但该增益仅适用于含金额/时序的账户交易图，少量标注校准的进一步增益不稳定。'
        '核心 insight：反诈 GNN 的瓶颈不在算法设计，而在“盲扫 vs 扩线”的任务设定与无监督方法的适用边界。学生写 300 字反思：'
        '“作为未来的警务技术使用者，你在什么情况下会信任这个系统？什么情况下不会？”')

    add_left_heading(doc, '5.4  考核方式设计', level=2)
    add_para(doc,
        '考核采用“过程+产物”双轨（表 4），评价重点不是指标最大化，而是学生对“系统能力与局限”的如实把握——既看到 AI 辅助研判的价值，也理解其在真实场景中的局限，'
        '这一取向本身即承载科学素养训练。')
    add_table(doc,
        [
            ['类别', '观测指标', '权重', '评价要点'],
            ['过程', '案情分析合理性', '15%', '能否识别案件关键实体与资金流向'],
            ['过程', '人机对比串并案', '15%', '能否分析人与系统各自的优势与盲区'],
            ['过程', '冻卡决策理由', '10%', '能否论证误冻/漏冻权衡的决策依据'],
            ['产物', '边界认知反思', '30%', '是否理解“合成≠真实、盲扫全败”'],
            ['产物', '伦理分析报告', '15%', '冻卡决策中的权责分析深度'],
            ['产物', '团队答辩', '15%', '表达与协作'],
        ],
        '表 4  实训评价量表（权重为设计建议）',
        col_widths=[0.7, 1.8, 0.6, 2.2])

    add_left_heading(doc, '5.5  课程思政与伦理教育', level=2)
    add_para(doc,
        '反诈主题天然承载思政育人价值：守护群众财产安全、法治意识、技术向善与总体国家安全观。参照网络空间安全专业课程思政的融入经验[13]，'
        '本实训将思政与伦理元素嵌入各 Lab 而非附加说教（表 5）：Lab1 强调数据合规与“数据不出域”，Lab2 树立“AI 辅助而非替代民警”的人机协同意识，'
        'Lab3 在冻卡决策中开展公民权利保护讨论，Lab4 培育诚实严谨的科学精神。实训要求学生就“研判能力的使用边界与数据合规”撰写伦理分析，将“能研”与“不可滥用”直接对应到实训环节。')
    add_table(doc,
        [
            ['实训 Lab', '思政/伦理点', '融入方式'],
            ['Lab1（案情分析）', '数据合规意识', '脱敏数据、“数据不出域”讨论'],
            ['Lab2（串并案）', '人机协同定位', 'AI 辅助研判 ≠ 替代民警'],
            ['Lab3（冻卡决策）', '公民权利保护', '误冻/漏冻权衡·角色扮演'],
            ['Lab4（边界反思）', '科学精神与诚信', '合成 ≠ 真实、不夸大能力'],
        ],
        '表 5  实训模块与思政/伦理点映射',
        col_widths=[1.4, 1.4, 2.5])

    add_left_heading(doc, '5.6  教学成效说明与实施预案', level=2)
    add_para(doc,
        '本案例设计尚未在正式课程中规模化实施。为保障教学成效评估的客观性，本节按“预注册”原则预先声明评价指标，待 2026 年秋季学期实施后按同样指标采集数据并如实报告，'
        '避免选择性报告带来的偏差。评价体系含六个维度：案情分析能力（课前/课后问卷）、人机协同意识（Lab2 报告盲评）、决策伦理分析深度（冻卡报告盲评 4 级量表）、'
        '诚实边界认知（反思中“合成≠真实”表述准确性）、系统可用性（SUS 量表）与课程推荐意愿（NPS），各指标在实施前声明、实施后按同样标准采集。')

    add_left_heading(doc, '5.7  教学优先的设计哲学', level=2)
    add_para(doc,
        '本案例以 4 节 Jupyter Notebook 实验课为核心载体，以研判决策而非模型分析为核心活动，基于三层考量。其一，GNN 作为工具而非教材：学生不需要理解 HAN 内部原理或计算 F1，'
        '而以办案人员视角使用系统、做出决策、反思边界。其二，预计算降低门槛：系统研判结果预烘焙，学生 Run All 即可，教学时间 100% 用于决策、对比与反思。'
        '其三，决策驱动认知：Lab1 认识案情 → Lab2 人机对比 → Lab3 冻卡决策 → Lab4 边界反思，认知负荷逐步提升。'
        'FraudLens 完整原型（docker-compose 全栈）保留为进阶入口，供有工程兴趣的学生在课程设计或毕设中深入，形成“基础层（4 Lab 研判决策）—进阶层（全栈原型工程）”的双层设计。')

    # Section 6
    add_centered_heading(doc, '6  结论与展望', level=1)
    add_para(doc,
        '针对公安院校反诈实战教学中学生难以动手操作真实研判系统、教师缺乏可直接落地的现成案例等问题，本文以可运行的反诈团伙研判原型系统 FraudLens 为载体，'
        '设计了一套面向学生能力成长的 AI 赋能实验教学案例，形成了“研—训—评”一体化的实训模式与四环节递进实验链。'
        '与聚焦通用编程或单门课程的现有案例相比，本案例的教学创新主要体现在三方面：')
    innovations = [
        '真实系统转化为可复现“研”对象：FraudLens 两级研判闭环与 LangGraph 反思工作流为学生提供了可运行、可逐段复现的真实研究对象；'
        '学生无需掌握 GNN 内部机理即可在 Jupyter Notebook 中一键复现能力基线，将“不可触碰的前沿系统”变为“可拆解的教学载体”。',
        '以研判决策为核心组织“训”的链条：四 Lab 从案情分析、人机协同串并案、冻卡伦理决策到能力边界反思逐级递进；'
        '学生在“人 vs AI”对比中建立人机协同意识，在阈值调参与角色扮演中体会技术—伦理权衡，在失败案例中形成科学诚实与算法边界认知。',
        '以过程+产物双轨“评”支撑育人成效：评价量表将“能否如实把握系统能力与局限”作为核心观测点，'
        '思政/伦理映射把数据合规、公民权利保护、科学诚信内嵌于实训环节，避免技术训练与价值引领“两张皮”。'
    ]
    for idx, txt in enumerate(innovations, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r = p.add_run(f'（{idx}）{txt}')
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r.font.size = Pt(12)
    add_para(doc,
        '通过本案例，学生可望在复杂、不确定的反诈研判场景中形成三项核心能力：AI 辅助研判决策能力——理解系统输出并判断“何时可信、何时不信”；'
        '人机协同能力——识别 AI 与人的互补盲区；工程可信与伦理责任能力——在冻卡决策等真实后果场景中做出有依据、有担当的判断。'
        '上述能力直接服务于公安院校培养“懂技术、会研判、守边界”的实战化人才目标。')
    add_para(doc,
        '局限与展望：本案例目前为教学设计论证阶段，尚未在正式课程中规模化实施，教学成效有待 2026 年秋季学期按预注册指标采集后报告。'
        '后续将与公安机关合作获取脱敏真实数据，进一步优化门控策略与跨任务迁移能力，持续完善可复现实验资源与评价工具。')

    # References
    add_centered_heading(doc, '参考文献', level=1)
    refs = [
        '李清勇，等. 私教还是枪手：基于通用大语言模型的计算机实践教学探索与反思[J]. 实验技术与管理，2024，41（5）：1-8.',
        '张金，宫晓利，高小鹏，等. 基于通用大语言模型的计算机系统创新实验设计[J]. 实验技术与管理，2024，41（10）：1-9.',
        '向尕，等. 新工科背景下“解决复杂工程问题”能力培养研究——以信息安全专业综合实习为例[J]. 软件导刊，2022，21（9）：211-218.',
        'DOU Y T, LIU Z W, SUN L, et al. Enhancing graph neural networks by a label propagation algorithm for fraud detection[C]//Proceedings of the ACM International Conference on Information & Knowledge Management. 2020: 2589-2592.',
        '蒋宗礼. 本科工程教育：聚焦学生解决复杂工程问题能力的培养[J]. 中国大学教学，2016（11）：27-30,84.',
        'LANGCHAIN. LangGraph: Building stateful multi-agent applications with LLMs[EB/OL]. (2024)[2026-07-31]. https://github.com/langchain-ai/langgraph.',
        'HAMILTON W L, YING R, LESKOVEC J. Inductive representation learning on large graphs[C]//Advances in Neural Information Processing Systems. 2017.',
        'WANG X, JI H, SHI C, et al. Heterogeneous graph attention network[C]//The World Wide Web Conference. 2019: 2022-2032.',
        'BAAI. BGE: BAAI general embedding[EB/OL]. (2023)[2026-07-31]. https://huggingface.co/BAAI/bge-large-zh-v1.5.',
        'SHINN N, CASSANO F, GOPINATH A, et al. Reflexion: Language agents with verbal reinforcement learning[C]//Advances in Neural Information Processing Systems. 2023.',
        'WEBER M, DOMENICONI G, CHEN J, et al. Anti-money laundering in Bitcoin: Experimenting with graph convolutional networks for financial forensics[J/OL]. arXiv:1908.02591, 2019.',
        'WEBER M, DOMENICONI G, CHEN J, et al. Scalable graph learning for anti-money laundering: A first look[C/OL]. arXiv:1812.00076, 2018.',
        '李剑. 网络空间安全专业研究生课程思政教育的探索与实践[J]. 信息安全研究，2024，10（2）：190-192.',
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'[{i}] {r}')
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(10.5)

    out_path = os.path.join(BASE, 'CPEC2026_draft.docx')
    doc.save(out_path)
    print('Saved', out_path)


if __name__ == '__main__':
    build()
