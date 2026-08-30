# -*- coding: utf-8 -*-
"""FraudLens 项目计划书 v2 —— 黑白学术排版版（正式书面语修订）。
标题体系：一级“一、”（黑体15）/ 二级“（一）”（黑体13）/ 三级“1.”（黑体12）；
正文列举用（1）（2）编号段；仅配图允许彩色；数据表为学术三线表；
图表标题一律置于图表上方（keep_with_next 防孤题），表注按学术惯例置于表下。
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SRC = r"E:\FraudLens\docs\拟定模板_最新.docx"
DST = r"E:\FraudLens\docs\项目计划书（已填写）.docx"
FIG = r"E:\FraudLens\docs\plan_figures_v2"

BLACK = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------- 基础工具
def set_font(run, cn="宋体", size=12, bold=False, color=BLACK, italic=False):
    run.font.name = cn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def first_line_indent(p, chars=2):
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = pPr.makeelement(qn("w:ind"), {})
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, v in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def three_line_table(table):
    """学术三线表：仅顶线/底线（1.5pt），无竖线无内部横线。"""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "12")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    for edge in ("left", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)


def cell_bottom_border(cell, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), "000000")
    borders.append(b)
    tcPr.append(borders)


def clear_cell(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    return cell.paragraphs[0]


def add_rich(p, text, size=12, cn="宋体", color=BLACK, bold_all=False):
    """支持 **加粗** 标记（黑色加粗，不用颜色强调）。"""
    tokens = re.split(r"(\*\*.+?\*\*)", text)
    for seg in tokens:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            run = p.add_run(seg[2:-2])
            set_font(run, cn, size, True, color)
        else:
            run = p.add_run(seg)
            set_font(run, cn, size, bold_all, color)


def row_cells(row):
    seen = []
    for c in row.cells:
        if c not in seen:
            seen.append(c)
    return seen


def set_cell_replaced(cell, old, new, size=10.5):
    orig = cell.text
    p = clear_cell(cell)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(orig.replace(old, new))
    set_font(run, "宋体", size, False)


def check_cell(row, old, new):
    cells = row_cells(row)
    set_cell_replaced(cells[-1], old, new)


# ---------------------------------------------------------------- 排版元素
def add_h1(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, "黑体", 15, True)


def add_h2(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, "黑体", 13, True)


def add_h3(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    first_line_indent(p, 2)
    run = p.add_run(text)
    set_font(run, "黑体", 12, True)


def add_body(cell, text, size=12):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    first_line_indent(p, 2)
    add_rich(p, text, size=size)


def img_wh_cm(name, max_w=14.2, max_h=18.5):
    with Image.open(os.path.join(FIG, name)) as im:
        w, h = im.size
    r = h / w
    w_cm, h_cm = max_w, max_w * r
    if h_cm > max_h:
        h_cm = max_h
        w_cm = h_cm / r
    return Cm(w_cm), Cm(h_cm)


_pending_cap = []


def _squash_after_tbl(tbl):
    """把 add_table 自动追加在表后的空段落压到近零高度，消除图/表块下方多余空白。"""
    nxt = tbl._tbl.getnext()
    if nxt is not None and nxt.tag == qn("w:p"):
        pPr = nxt.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        sp.set(qn("w:line"), "20")
        sp.set(qn("w:lineRule"), "exact")
        pPr.append(sp)
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            pPr.append(rPr)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "2")
        rPr.append(sz)


def add_fig(cell, name, scale=1.0):
    """图块：无边框单格表承载图题+图，行级 cantSplit 禁止跨页拆分（图题在上、图在下，同页）。"""
    cap = _pending_cap.pop(0) if _pending_cap else None
    tbl = cell.add_table(rows=1, cols=1)
    _squash_after_tbl(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))
    c = tbl.cell(0, 0)
    p0 = c.paragraphs[0]
    if cap:
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(8)
        p0.paragraph_format.space_after = Pt(2)
        run = p0.add_run(cap)
        set_font(run, "黑体", 10.5, False)
    else:
        p0._p.getparent().remove(p0._p)
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    w, h = img_wh_cm(name)
    run = p.add_run()
    run.add_picture(os.path.join(FIG, name),
                    width=Cm(w.cm * scale), height=Cm(h.cm * scale))


def add_cap(cell, text):
    """图题：暂存文本，由随后的 add_fig 合并进同一图块（置图上方）。"""
    _pending_cap.append(text)


def add_data_table(cell, header, rows, widths, bold_cells=None, title=None, note=None):
    """学术三线表：表题在上、注释在下；无底纹、无竖线，仅黑色加粗强调最优值。"""
    if title:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title)
        set_font(run, "黑体", 11.5, True)
    t = cell.add_table(len(rows) + 1, len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _squash_after_tbl(t)
    hdr_trPr = t.rows[0]._tr.get_or_add_trPr()
    hdr_trPr.append(OxmlElement("w:tblHeader"))
    three_line_table(t)
    bold_cells = bold_cells or set()
    for j, htxt in enumerate(header):
        c = t.cell(0, j)
        set_cell_margins(c, 40, 40, 60, 60)
        cell_bottom_border(c, "6")
        p = clear_cell(c)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(htxt)
        set_font(run, "黑体", 9.5, True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i + 1, j)
            set_cell_margins(c, 30, 30, 30, 30)
            p = clear_cell(c)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            is_b = (i, j) in bold_cells
            run = p.add_run(val)
            set_font(run, "宋体", 9, is_b)
    for j, wd in enumerate(widths):
        for i in range(len(rows) + 1):
            t.cell(i, j).width = Cm(wd)
    if note:
        # 表注包进无边框 cantSplit 嵌套表：Word 对单元格内段落不执行 keepLines，
        # 只有表格行 cantSplit 能保证表注整段不跨页拆行。
        nt = cell.add_table(1, 1)
        nt.alignment = WD_TABLE_ALIGNMENT.CENTER
        _squash_after_tbl(nt)
        ntblPr = nt._tbl.tblPr
        nb = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            e.set(qn("w:space"), "0")
            nb.append(e)
        ntblPr.append(nb)
        ntrPr = nt.rows[0]._tr.get_or_add_trPr()
        ntrPr.append(OxmlElement("w:cantSplit"))
        nc = nt.cell(0, 0)
        nc.width = Cm(sum(widths))
        set_cell_margins(nc, 0, 0, 0, 0)
        p = clear_cell(nc)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(8)
        add_rich(p, note, size=9)
    else:
        cell.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------- 文档构建
doc = Document(SRC)
table = doc.tables[0]

r0 = row_cells(table.rows[0])
set_cell_simple_text = r0[1]
_p = clear_cell(r0[0])
for _r in list(_p.runs):
    _r._element.getparent().remove(_r._element)
_run = _p.add_run("项目名称")
set_font(_run, "宋体", 10.5, False)
_p = clear_cell(r0[1])
for _r in list(_p.runs):
    _r._element.getparent().remove(_r._element)
_run = _p.add_run("FraudLens——面向反诈一线的多智能体与图神经网络团伙发现系统")
set_font(_run, "宋体", 10.5, True)
check_cell(table.rows[1], "□创意组", "☑创意组")
check_cell(table.rows[2], "□主赛道", "☑主赛道")
check_cell(table.rows[3], "□“人工智能+”项目", "☑“人工智能+”项目")
check_cell(table.rows[4], "□创意计划阶段，尚未注册公司", "☑创意计划阶段，尚未注册公司")

cell = row_cells(table.rows[19])[1]
anchor = clear_cell(cell)
for r in list(anchor.runs):
    r._element.getparent().remove(r._element)

# ================================================================ 执行摘要
add_h1(cell, "执行摘要")
add_body(cell, "**项目概述**：FraudLens 是面向公安反诈一线的智能辅助研判系统。办案民警上传批量诈骗线索（话术文本、案件表格、聊天截图）后，系统自动将分散个案关联为关系网络，基于图神经网络从中识别潜在诈骗团伙，并为每一条串并建议输出完整证据链，将传统依赖人工经验与肉眼比对的串并案研判，升级为可解释、可复核、新民警即可上手的智能化流程。")
add_cap(cell, "图 1　项目一页通：核心指标、四大板块与评分点对照")
add_fig(cell, "fig1_summary.png")
add_body(cell, "**核心优势与实证基础**：本项目的全部效果结论均来自可复现的真实实验，而非停留在仿真数据上的算法论证。在 6 种警情预设 × 5 种基线方法 × 3 随机种子的全量对照中，最贴近真实警情的重噪场景下本方案 pairwise F1 达 **0.9487±0.0410**，三个随机种子全部优于基线且方差最小；系统前后端已完成开发并可运行，支持本地化部署、数据不出城；团队为警校信息安全专业本科生，选题源于反诈宣讲中获取的一线实证需求，反诈中心实地走访验证正在进行。")

# ================================================================ 一、背景与问题
add_h1(cell, "一、项目背景与问题界定")
add_h2(cell, "（一）电信网络诈骗已成为发案最多、损失最大的犯罪类型")
add_body(cell, "公安部 2026 年 1 月新闻发布会公布：**2025 年全国侦破电信网络诈骗案件 25.8 万起**，抓获头目骨干 542 名，拦截诈骗电话 36 亿次、短信 33 亿条，封堵涉诈域名 816.2 万个，紧急止付涉诈资金 **2170.7 亿元**，见面劝阻 674.7 万人次。最高检《刑事检察工作白皮书（2025）》显示：起诉电信网络诈骗、帮助信息网络犯罪活动、掩饰隐瞒犯罪所得三类犯罪人数已占全部刑事起诉的 **14.3%**，即每 7 名被起诉的刑事被告人中约有 1 名与电诈及其关联犯罪有关。打击治理电信网络诈骗已从刑事侦查的重要分支，上升为平安建设的主战场。")
add_body(cell, "与高压打击形成对照的，是一线研判方式的相对滞后。当前诈骗分子呈现团伙化、跨地域、话术模板化的作案特征：一个团伙同时操控多部手机号与多个银行账户，案件分散报送至不同地区公安机关。**个案呈点状分布，关联后方显团伙全貌**——但从“点”到“团”的串并研判环节，目前仍主要依靠民警人工完成。")
add_h2(cell, "（二）一线串并案研判的四大痛点")
add_body(cell, "（1）**串并依赖人工**：同一团伙流窜作案被拆分为数十条孤立线索，人工比对手机号、翻查账户流水、凭经验关联案件，漏并、误并情况时有发生；")
add_body(cell, "（2）**话术迭代加速**：诈骗话术模板化生成、高频微调，关键词规则库存在天然滞后，难以跟上话术演变速度；")
add_body(cell, "（3）**数据质量受限**：报警笔录中账户信息缺失、号码打码、金额口径不一，纯规则与纯文本相似度方法在真实警情数据上均难以奏效（本项目实验证实：纯资金链规则在重噪场景仅 0.82，纯话术聚类仅 0.50，详见第四部分）；")
add_body(cell, "（4）**增量警情持续涌入**：案件每日新增，传统批量重算模式需攒批处理、全量重跑，计算与时间成本难以适应基层“随来随录”的实战节奏。")
add_h2(cell, "（三）政策环境与技术窗口期")
add_body(cell, "《反电信网络诈骗法》自 2022 年 12 月施行，明确要求各部门协同治理、加强技术反制能力；浙江等地已发布“人工智能+公安”行动计划。现有大厂反诈平台以**事前预警拦截**为主（面向海量通信与金融数据），而案发后“从已报案线索中挖掘团伙”的**串并研判环节**，仍是县区级反诈中心依赖人力投入、长期缺乏专用工具支撑的薄弱环节。大模型与图神经网络技术的成熟，使该环节首次具备低成本自动化改造的可行性。")
add_h2(cell, "（四）问题来源与调研基础")
add_body(cell, "本项目选题源于一线实战需求，而非文献推演。2025 年寒假，团队成员在家乡反诈宣讲中直接接触到群众受骗案例，并首次听基层民警反映“线索串不起来、团伙看不全”的实际困难。此后团队系统研读了公安部历年通报、最高检白皮书与公开判例中的作案模式描述，据此设计了 6 种贴近实战的警情数据预设（干净/轻噪/重噪/规模化/话术碰撞/碰撞+噪声）。**下一步（正在进行）**：依托警校师生渠道对接县分局反诈中心，开展实地走访调研，以真实工单回访与试用反馈校准系统，调研结论将补充进后续版本材料与论文。")

# ================================================================ 二、方案
add_h1(cell, "二、系统总体方案与架构设计")
add_body(cell, "FraudLens 以“**线索输入、团伙输出**”为目标设计为四层架构：输入层兼容话术文本、CSV 批量线索、聊天截图（OCR 自动取文）三类一线真实数据形态；多智能体研判层基于 LangGraph 状态图编排 5 个专业 Agent 流水线作业；图学习层为团伙发现引擎；输出层提供串并案分组、团伙画像、研判报告与资金流向可视化。")
add_cap(cell, "图 2　FraudLens 总体架构：数据全程本地处理，不出城")
add_fig(cell, "fig2_arch.png")
add_h2(cell, "（一）多智能体协同研判与反思质检机制")
add_body(cell, "线索接入 Agent 完成清洗、要素抽取与字段归一；话术分析 Agent 识别话术模板与作案阶段（引流/实施/洗钱）；资金链 Agent 追踪账户共享、转账链条与取现模式；串并建议 Agent 综合三路信息给出“串并假设 + 证据链”；**反思质检 Agent** 对三路结论交叉验证，不达标即驳回重判，直至通过或触顶降级为“人工复核”。区别于简单的提示词串联式伪多智能体架构，本系统的反思闭环在 LangGraph 状态图上是真实的回边结构，驳回、重试、降级全过程可追溯。")
add_h2(cell, "（二）异构图谱建模：以证据关联定义案件关联")
add_body(cell, "系统将案件、受害人、手机号、收款账户、话术模板、城市六类对象构建为**异构图**，案件之间通过 5 条元路径产生关联：共享账户（硬边，强证据）、共享号码（硬边）、话术语义相似（BGE 中文语义向量余弦构成的**软权重边**）、同城、同受害人。图学习的价值在于：两个案件即使没有任何单一字段完全相同，只要在图中被同一批证据节点关联，模型即可将其识别为同一团伙的组成部分。")
add_cap(cell, "图 3　诈骗团伙异构构图示意：灰点案件“话术相似但资金零重合”，被双信号门控拒绝并入，不误并")
add_fig(cell, "fig3_graph.png")

# ================================================================ 三、创新
add_h1(cell, "三、核心创新：面向低质量实战数据的团伙发现方法")
add_body(cell, "真实警情数据的质量受限特征，是本项目算法设计的出发点，也是本项目与通用知识图谱方案的技术分水岭。核心技术包括三项：**共识伪标签半监督、客观置信度门控、增量匹配**。")
add_h2(cell, "（一）创新一：共识伪标签半监督——规则提供锚点、图网络完成泛化（专利核心主张）")
add_body(cell, "团伙标注的获取是首要难题：请民警标注成本过高，纯无监督方法在噪声数据上性能有限（自监督 GNN 在重噪数据上仅 0.565，见第四部分）。本项目的方案：让**两条相互独立的证据通道交叉验证**——资金通道采用 Louvain 社区发现为案件赋予资金链标签，话术通道采用 BGE 语义向量与 KMeans 为案件赋予话术标签；仅当**两通道结论严格一致**时，对应案件才被采纳为“共识锚点”（实测锚点纯度 100%）。再以锚点作为伪标签监督微调 HAN 图注意力网络：锚点案件直接采用标签，非锚点案件采用微调后模型预测。该方法具有清晰的数学性质：**整体识别率的下界等于锚点质量**，即系统已确认的结论具备可证明的可靠性保障。")
add_cap(cell, "图 4　共识伪标签半监督全流程：双通道独立取证 → 一致才可信 → 锚点监督 GNN")
add_fig(cell, "fig4_semi.png")
add_h2(cell, "（二）创新二：客观置信度门控——具备拒绝输出能力的审慎判定机制")
add_h3(cell, "1. 三级拒绝规则")
add_body(cell, "孤案不成簇不予采信（防止“自我确认”式误并）；簇规模超过全库 25% 判定为社区发现污染，整类弃用（200 案规模实测曾出现 90 案巨簇，若不设防将导致全库误并为单一团伙）；可用锚点不足 2 簇时，整个半监督环节**拒绝输出结论**，回退人工研判提示。")
add_h3(cell, "2. 自适应 k 重估机制")
add_body(cell, "在三级拒绝规则之上，本项目设计了**自适应 k 重估机制**：当默认聚类粒度产生巨簇时，算法沿更细粒度自动扫描，直至不再触发巨簇弃用——触发者是数据自身特征，而非以测试集指标反推参数。全量三种子复测结果：轻噪场景本方案由 0.692 提升至 **0.756**，规模化场景由 0.466 提升至 **0.656（+0.19）**，干净基线由 0.812 提升至 **0.913**；重噪与话术碰撞场景无巨簇、结果不变；碰撞+噪声场景此前因巨簇弃用导致一个种子拒绝输出（均值样本数记为 2），自适应后恢复输出，均值如实按 3 个种子统计。")
add_cap(cell, "图 5　三道保险：三级拒绝规则、自适应 k 重估、实测拒绝输出案例")
add_fig(cell, "fig5_gating.png")
add_h3(cell, "3. 拒绝输出能力的实战意义")
add_body(cell, "警务辅助决策的基本要求是：宁可不出结论，不出低置信度结论。早期全量实验中，P5 场景某一随机种子曾因“共识锚点不足”真实触发整体拒绝输出——当时如实标注该组均值样本数，未作掩盖或强行输出；自适应 k 重估后该种子恢复输出（单值偏低仍如实保留）。**门控语义未被扭曲，只是扩大了“可以有把握输出结论”的适用范围**。在办案责任制背景下，不具备拒绝输出能力的辅助研判系统难以投入实战应用。")
add_h2(cell, "（三）创新三：增量匹配——面向持续增量警情的流式团伙归并")
add_body(cell, "针对案件每日新增的实战节奏，系统为每个已发现团伙维护画像（成员账户池与 BGE 话术质心），新线索进入时执行轻量匹配：**仅当资金共享与话术语义双信号一致时并入团伙，否则暂缓处理、待证据累积后复核**——延续审慎合并的门控原则。200 案规模实测：增量模式识别率 0.6089，优于全量 Louvain 模式的 0.5263（**+0.083**）——全量社区发现在大规模数据下存在过合并倾向，流式逐案归并反而能更精细地划分团伙边界。")

# ================================================================ 四、实验
add_h1(cell, "四、实验验证与结果分析")
add_h2(cell, "（一）实验设计与可复现性")
add_body(cell, "本项目自建了可复现的实验框架：**固定随机种子与 sha256 确定性哈希**保证任何人在任何机器重跑均得到完全一致的结果；6 种警情预设模拟从“干净台账”到“账户缺失、话术变体、规模化、话术模板互相碰撞”的真实难度梯度；与 5 类基线方法对照（KMeans 纯聚类、纯资金链规则、自监督 GNN、纯 GNN 输出、本方案半监督混合策略）。评价指标采用 pairwise F1，衡量系统分组结果与真实团伙的吻合程度（取值 0~1，越高越准确），3 个随机种子取均值 ± 标准差。")
add_h2(cell, "（二）全量实验结果")
add_cap(cell, "图 6　全量实验总览：主战场条形图 + 6 预设全景表")
add_fig(cell, "fig6_experiment.png", scale=0.87)
add_data_table(
    cell,
    ["警情预设", "KMeans", "资金链规则", "自监督GNN", "纯GNN", "本方案", "场景最优"],
    [
        ["P0 干净基线", "0.853±0.021", "0.833±0.129", "0.922±0.057", "0.683±0.130", "0.913±0.030", "自监督GNN"],
        ["P1 轻噪", "0.704±0.130", "0.819±0.056", "0.566±0.054", "0.546±0.109", "0.756±0.127", "资金链规则"],
        ["P2 重噪·主场景", "0.499±0.064", "0.821±0.118", "0.565±0.069", "0.893±0.078", "0.949±0.041", "本方案"],
        ["P3 200案规模", "0.540±0.089", "0.447±0.097", "0.536±0.118", "0.642±0.073", "0.656±0.069", "本方案"],
        ["P4 话术碰撞", "0.846±0.093", "1.000±0.000", "0.890±0.156", "0.629±0.226", "0.981±0.027", "资金链规则"],
        ["P5 碰撞+噪声", "0.712±0.229", "0.833±0.129", "0.623±0.073", "0.737±0.133", "0.797±0.174", "资金链规则"],
    ],
    [2.6, 1.95, 1.95, 1.95, 1.95, 1.95, 1.95],
    bold_cells={(0, 3), (1, 2), (2, 5), (3, 5), (4, 2), (5, 2)},
    title="表 1　全量实验：6 种警情预设 × 5 方法 × 3 随机种子（pairwise F1，mean±std）",
    note="注：数值为自适应 k 门控优化后的全量三种子（42/7/2024）实验结果，固定种子可复现；加粗为该行最高值。P5 种子 7 单值 0.555 偏低，详见局限性（4）。",
)
add_body(cell, "**主战场结论（P2 重噪，最贴近真实警情：账户缺失 + 话术变体）**：本方案 0.949±0.041 为该场景最优，三个随机种子全部优于基线且方差最小——共识伪标签半监督的有效性经多随机种子验证，排除了偶然波动因素。**规模化场景（P3）本方案 0.656 亦为五种方法中最优**（绝对水平偏低的原因见下文局限性说明）；干净基线（P0）自监督 GNN 0.922 最优、本方案 0.913 紧随其后。整体规律明确：**数据噪声越强、规模越大，本方案相对优势越显著**——这正是反诈一线的真实数据处境。")
add_h2(cell, "（三）局限性主动呈报")
add_body(cell, "（1）干净数据（P0）上锚点过分裂的影响虽经自适应 k 缓解仍未完全消除，本方案 0.913 略低于自监督 GNN 0.922，已作为论文局限性讨论，不以针对性调参掩盖；")
add_body(cell, "（2）规模化场景（P3）绝对水平仍偏低（本方案 0.656），根源是 Louvain 算法的分辨率极限（文献已知问题），解决路径为增量匹配架构（见前文创新三，实测优于全量 +0.083），而非继续调参；")
add_body(cell, "（3）轻噪（P1）与碰撞+噪声（P5）场景资金链规则仍为最优（0.819/0.833），本方案 0.756/0.797 居次——规则方法在轻噪条件下本身较强，差距已收窄至 0.06 左右，半监督方法的价值在重噪、规模化场景更为突出；")
add_body(cell, "（4）P5 场景种子 7 单值 0.555 偏低（极端信号下共识锚点弱），如实保留、不做剔除；")
add_body(cell, "（5）早期开发中曾发现“hash 向量冒充语义嵌入”“进程间哈希随机化”两个隐蔽缺陷，当日 15:00 前的全部结论已废弃重跑——**该问题发现与修正过程本身即为工程可信度的佐证**。")

# ================================================================ 五、产业
add_h1(cell, "五、产业认知与竞争定位")
add_h2(cell, "（一）产业格局与主要参与者")
add_body(cell, "参与竞争首先须客观认识现有格局。美亚柏科等上市安全厂商的反诈预警平台已部署近 60 个公安机关，累计预警 5665 万余条、避免损失 203.9 亿元，其优势在于**资金规模大、平台级交付、部省市贯通**，主攻海量数据的事前预警拦截。同盾等金融风控厂商的优势在交易事前拦截，不面向公安案发后的“多案归团、证据组织”流程。通用知识图谱平台面向数据工程人员，需专人建模调参，并非面向办案民警的直接可用工具。")
add_cap(cell, "图 7　产业四象限定位与差异化切入路径")
add_fig(cell, "fig7_market.png")
add_h2(cell, "（二）市场卡位、竞争壁垒与切入路径")
add_body(cell, "FraudLens 定位于四象限的右下格：**县区级反诈中心日常串并研判**——头部企业未重点覆盖、基层刚需明确细分市场。三条竞争壁垒：① **零标注**——共识伪标签半监督使反诈中心无需配置算法工程师即可使用；② **可解释**——每条串并建议附账户、号码、话术证据链，支撑立案说明，直接对接办案责任制；③ **增量+门控**——新警情流式归并不重算全图，置信度不足即拒绝输出。切入路径三步走：本校公安专业师生渠道 → 县分局反诈中心免费试点（换取真实数据与场景反馈）→ 地市反诈平台“研判插件”位 → 依托实证成果申报公安部科技创新项目。")

# ================================================================ 六、商业模式
add_h1(cell, "六、商业模式与落地路径")
add_h2(cell, "（一）收入模式与成长逻辑")
add_body(cell, "项目处于创意阶段，此处不作营收预测，重点说明**收入来源、市场容量与成长逻辑**。收入模式：前期以免费试点换取真实数据与联合研究成果（论文、专利、示范案例）；中期面向地市公安局提供“平台研判插件 + 本地化部署 + 模型定制”的项目制服务（公安信息化单个项目普遍为数十万至数百万元级）；长期以“团伙画像订阅 + 跨域串并协作网络”形成持续性收入——串并案天然跨区域，**用户规模与图谱覆盖度、识别精度形成正向循环**，具备数据网络效应。")
add_h2(cell, "（二）成本与合规")
add_body(cell, "成本与合规：系统单机 CPU 即可运行全流程（本项目实验环境即为纯 CPU 环境），县区级可复用现有办案硬件，边际成本接近于零；**案件数据全程本地处理、不出城**，BGE 语义模型本地部署，从架构层面满足《反电信网络诈骗法》与公安数据管理规范对个人信息的刚性要求——这是云端方案在县区一线难以落地的合规约束，恰是本项目的立足点。")

# ================================================================ 七、团队
add_h1(cell, "七、团队构成与核心能力")
add_body(cell, "团队由**湖北警官学院信息安全专业本科生**组成。“警校背景与信息安全专业能力”的复合结构是本项目最难复制的优势：算法团队具备阅读图神经网络前沿论文的能力，同时熟悉“串并”“引流”“取现模式”等业务术语；全栈开发团队独立完成前后端、可视化与本地部署；调研验证团队负责一线需求访谈、实验设计与数据治理；统筹文档团队负责申报材料与知识产权布局。")
add_body(cell, "分工以项目里程碑为轴而非以人为轴：算法创新（共识半监督/门控/增量）、系统实现（LangGraph 流水线/构图/可视化）、实证闭环（实验框架/走访调研）、成果输出（论文/专利/申报书）四条线并行推进、周度对齐。依托警校渠道，团队成员具备接触一线反诈业务的天然入口；依托信息安全专业训练，团队具备将业务语言转化为算法方案的完整能力。**警校背景构成需求入口优势，全栈能力保障工程交付。**")

# ================================================================ 八、实施与成长
add_h1(cell, "八、实施计划与个人成长")
add_h2(cell, "（一）实施路线图")
add_cap(cell, "图 8　实施路线图：已完成 62%，竞赛验证期进行中")
add_fig(cell, "fig8_roadmap.png")
add_body(cell, "研发攻坚期（2025.09–2026.08，已完成）：从家乡反诈宣讲立项，到多智能体流水线与异构 GNN 首版跑通，再到三种子全量实验定型、脱敏数据合成管线与前端报告系统交付——**每一步均有代码提交与实验日志可查**。竞赛验证期（2026.09–2027.03，进行中）：反诈中心实地走访调研补上实证短板；中文论文《基于多智能体与图神经网络的反诈团伙发现》投稿；方法类发明专利 1–2 件提交（先申请占优先权日，再发表论文，保护新颖性）；申报书、演示视频、答辩 Demo 打磨。落地试点期（2027 起，规划）：县分局反诈中心免费试用、真实工单回访迭代；向涉赌、跨境网络赌博等团伙化警情做方法迁移验证（同一套共识半监督框架不需要新的标注体系）。")
add_h2(cell, "（二）能力成长与收获")
add_body(cell, "项目开展以来，团队能力实现了从“对公开数据集跑分”到“为真实社会问题**定义数据、设计实验、识别调参陷阱、对结论诚实**”的跨越。团队在项目中系统掌握了图神经网络、半监督学习、多智能体编排与全栈工程能力，更完成了从“学生作业思维”到“工程与科研并重”的转变——发现 hash 向量冒充语义嵌入的静默缺陷后，团队废弃一整天产出的全部结论、修复两个底层缺陷、全量重跑。这一决策过程，比任何实验分数都更能说明团队对“材料真实”要求的理解。")
add_h2(cell, "（三）价值体悟：把技术用在人民需要处")
add_body(cell, "作为警校生，团队对电诈之害有切近的认知：一通电话可能毁掉一个家庭的养老钱，一线民警需要在数百条线索中人工排查团伙。FraudLens 的每一次迭代都围绕同一标准——**这条建议民警是否敢用、用了能否经得起法庭质证**。把人工智能的研究写入“天下无诈”的平安中国建设，把个人成长融入守护人民群众财产安全的时代命题，这是本项目全部工作的意义所在。")

# ================================================================ 脚注
pf = cell.add_paragraph()
pf.paragraph_format.space_before = Pt(10)
run = pf.add_run("数据来源与可复现声明：本文全部实验数字来自 FraudLens 项目 2026 年 8 月全量三种子实验（固定种子 + sha256 确定性哈希，任何人可在 CPU 环境复现）；宏观数据来自公安部 2026 年 1 月新闻发布会、最高检《刑事检察工作白皮书（2025）》、证券时报对美亚柏科的公开报道，均可溯源。")
set_font(run, "宋体", 9, False)

doc.save(DST)
print("SAVED:", DST)
