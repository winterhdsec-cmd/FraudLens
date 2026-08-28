"""
File upload/extract/OCR/import routes.
Smart routing: text→tools, complex images→multimodal vision model
"""
import os
import io
import tempfile
import re

from fastapi import APIRouter, UploadFile, File, Query, Depends
from fastapi.responses import JSONResponse

from .deps import get_current_user
from core.llm_client import wrap_messages  # G2 脱敏

router = APIRouter(prefix='/api', tags=['文件'])

IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.webp')
TEXT_EXTENSIONS = ('.txt', '.csv')
DOC_EXTENSIONS = ('.docx',)
PDF_EXTENSIONS = ('.pdf',)

ALL_TEXT_EXTENSIONS = TEXT_EXTENSIONS + DOC_EXTENSIONS + PDF_EXTENSIONS

_llm_client = None
_llm_model = None


def _store_raw_to_minio(filename: str, content: bytes):
    """A4.1 (#C13): 原始上传件存 minio（证据留痕，数据不出域）。

    minio 不可用（SDK 缺失 / 服务未起）时返回 (None, None)，调用方退回本地，不报错。
    """
    try:
        from core.object_store import put_object, is_enabled
    except Exception:
        return None, None
    if not is_enabled() or not content:
        return None, None
    import hashlib
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    h = hashlib.md5(content).hexdigest()[:8]
    key = f"raw/{ts}_{h}_{filename}"
    # 存成功后返回后端代理下载地址（数据不出域：外部经 /api/object/{key} 取，避免 presigned 暴露 minio host 致签名失效）
    if put_object(key, content, content_type="application/octet-stream"):
        return key, f"/api/object/{key}"
    return None, None


def _get_llm_client():
    """获取 LLM 客户端（G2 统一网关；关闭/缺密钥返回 (None, None)）"""
    global _llm_client, _llm_model
    if _llm_client is not None:
        return _llm_client, _llm_model
    from core.llm_client import get_llm_client
    client = get_llm_client(sync=True)
    _llm_model = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
    _llm_client = client  # 可能为 None（关闭/缺密钥）
    return _llm_client, _llm_model


def _clean_chat_text(raw_text: str) -> str:
    """
    AI 后处理：对 OCR/视觉提取的文字进行清洗，
    区分嫌疑人（骗子）与受害人发言（聊天记录截图场景）。
    如果 LLM 不可用则返回原始文本。
    """
    if not raw_text or not raw_text.strip():
        return raw_text

    client, model = _get_llm_client()
    if not client:
        return raw_text

    prompt = (
        "你是一位反诈分析助手。以下是从图片中提取的原始文字（可能包含 OCR 识别错误），"
        "请进行以下处理：\n\n"
        "1. **清洗与纠错**：修正 OCR 识别错误，恢复正确的标点和分段\n"
        "2. **角色区分**：如果内容看起来是聊天记录或对话，请区分发言者角色：\n"
        "   - **[嫌疑人/骗子]**：冒充客服、公检法等身份，诱导转账、索要验证码等\n"
        "   - **[受害人]**：被诱导的受害人发言\n"
        "3. **内容归类**：按以下类别整理信息：\n"
        "   - 📞 通话/聊天内容（区分角色）\n"
        "   - 💳 资金信息（金额、账号、转账记录）\n"
        "   - 👤 人员信息（姓名、电话、身份证号）\n"
        "   - 📅 时间线信息\n"
        "4. **结构化输出**：如果内容较少，保留原始格式但添加角色标注\n\n"
        "请输出处理后的文本，保持原意不变，不要编造不存在的信息。\n\n"
        "原始文字：\n"
        f"{raw_text}"
    )

    try:
        from tools.response import logger
        response = client.chat.completions.create(
            model=model,
            messages=wrap_messages([{"role": "user", "content": prompt}]),
            temperature=0.1,
            max_tokens=2048,
            timeout=30
        )
        cleaned = response.choices[0].message.content
        if cleaned and cleaned.strip():
            logger.info(f"[TextClean] AI 文本清洗完成，{len(raw_text)}→{len(cleaned)} 字符")
            return cleaned
    except Exception as e:
        from tools.response import logger
        logger.warning(f"[TextClean] AI 清洗失败（{e}），返回原始文本")

    return raw_text


def _extract_docx(content: bytes) -> str:
    """提取 DOCX 段落 + 表格数据"""
    from docx import Document
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise ValueError(f"无法解析Word文件: {e}")
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(' | '.join(cells))
        parts.append(f'\n[表格 {i + 1}]\n' + '\n'.join(rows))
    return '\n'.join(parts)


def _extract_pdf_text(content: bytes) -> tuple:
    """
    提取 PDF 文字
    返回: (text, scanned)
    scanned=True 表示无文字层，需走 OCR/视觉通道
    """
    import pypdfium2 as pdfium
    pdf_doc = pdfium.PdfDocument(content)
    pages = []
    for i in range(len(pdf_doc)):
        page = pdf_doc[i]
        tp = page.get_textpage()
        text_page = tp.get_text_range()
        if text_page:
            pages.append(text_page.strip())
        tp.close()
    pdf_doc.close()
    total_text = '\n'.join(pages).strip()
    is_scanned = len(total_text) < 20 and len(pages) > 0
    return total_text, is_scanned


def _extract_pdf_content(file_bytes: bytes) -> str:
    """
    智能 PDF 提取: 有文字层直接提取, 无文字层转图片 OCR
    """
    text, is_scanned = _extract_pdf_text(file_bytes)
    if not is_scanned:
        return text
    from tools.ocr import ocr_image, ocr_image_file
    pages = _pdf_to_images(file_bytes)
    parts = []
    for i, img_bytes in enumerate(pages):
        ocr_result = ocr_image(img_bytes)
        if ocr_result.strip():
            parts.append(f'[第{i + 1}页]\n{ocr_result}')
    return '\n\n'.join(parts) if parts else ''


def _pdf_to_images(file_bytes: bytes) -> list:
    """PDF 每页转为 PNG 字节流列表"""
    import pypdfium2 as pdfium
    from PIL import Image
    pdf_doc = pdfium.PdfDocument(file_bytes)
    images = []
    for i in range(len(pdf_doc)):
        page = pdf_doc[i]
        bitmap = page.render(scale=2)
        pil_image = bitmap.to_pil()
        buf = io.BytesIO()
        pil_image.save(buf, format='PNG')
        images.append(buf.getvalue())
    pdf_doc.close()
    return images


# ──────────────────────────────────────────────
#  端点1: 传统文字提取（纯文字文档）
# ──────────────────────────────────────────────
@router.post('/extract-text')
async def api_extract_text(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    try:
        filename = file.filename or ''
        ext = os.path.splitext(filename)[1].lower()
        content = await file.read()
        
        # 调试：记录文件信息
        print(f"[DEBUG] 上传文件: {filename}, 扩展名: {ext}, 大小: {len(content)} bytes")
        
        text = ''
        source = 'direct'
        if ext in TEXT_EXTENSIONS:
            text = content.decode('utf-8', errors='ignore')
        elif ext == '.docx':
            print(f"[DEBUG] 开始解析DOCX文件...")
            text = _extract_docx(content)
            print(f"[DEBUG] DOCX解析成功，提取文本长度: {len(text)}")
        elif ext == '.pdf':
            text = _extract_pdf_content(content)
            source = 'pdf_ocr' if len(text) > 0 else 'direct'
        else:
            return JSONResponse(status_code=400, content={
                "success": False, "error": f"不支持的文件格式: {ext}"
            })
        object_key, object_url = _store_raw_to_minio(filename, content)
        return {"success": True, "text": text, "filename": filename, "source": source,
                "object_key": object_key, "object_url": object_url}
    except Exception as e:
        print(f"[ERROR] 文件处理失败: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


# ──────────────────────────────────────────────
#  端点2: OCR 文字识别（纯图片）
# ──────────────────────────────────────────────
@router.post('/ocr')
async def api_ocr_image(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    try:
        filename = file.filename or ''
        ext = os.path.splitext(filename)[1].lower()
        if ext not in IMAGE_EXTENSIONS:
            return JSONResponse(status_code=400, content={
                "success": False, "error": f"OCR 仅支持图片格式, 不支持: {ext}"
            })
        from tools.ocr import ocr_image
        content = await file.read()
        text = ocr_image(content)
        object_key, object_url = _store_raw_to_minio(filename, content)
        return {"success": True, "text": text, "filename": file.filename, "source": "ocr",
                "object_key": object_key, "object_url": object_url}
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


# ──────────────────────────────────────────────
#  端点3: 智能文件分析（核心改进 - 混合路由）
# ──────────────────────────────────────────────
@router.post('/analyze-file')
async def api_analyze_file(
    file: UploadFile = File(...),
    mode: str = Query('auto', description='auto: 自动选择, ocr: 强制OCR, vision: 强制多模态'),
    clean: str = Query('auto', description='auto: 图片结果自动AI清洗, off: 不清洗'),
    current_user: dict = Depends(get_current_user)
):
    """
    智能文件分析端点 - 自动选择最佳处理路径

    路由规则:
    - TXT/CSV/DOCX → 传统工具提取文字
    - PDF → 先检测是否有文字层 → 有则直接提取，无则转图片
    - 图片 → 根据复杂度自动选择:
        - 简单图片(白底文字) → EasyOCR (快)
        - 复杂图片(截图/表格) → DeepSeek VL2 多模态 (理解力强)
    - 强制 mode=ocr → 全部走 OCR
    - 强制 mode=vision → 图片走多模态
    - clean=auto → 图片结果自动 AI 清洗（区分嫌疑人/受害人）
    """
    try:
        filename = file.filename or ''
        ext = os.path.splitext(filename)[1].lower()
        content = await file.read()

        # ── 纯文字文档 ──
        if ext in TEXT_EXTENSIONS:
            text = content.decode('utf-8', errors='ignore')
            return {"success": True, "text": text, "filename": filename, "method": "direct", "route": "text"}

        if ext == '.docx':
            text = _extract_docx(content)
            return {"success": True, "text": text, "filename": filename, "method": "docx", "route": "text"}

        # ── PDF ──
        if ext == '.pdf':
            text, is_scanned = _extract_pdf_text(content)
            if not is_scanned:
                return {"success": True, "text": text, "filename": filename, "method": "pdf_text", "route": "text"}
            pages = _pdf_to_images(content)
            parts = []
            for i, img_bytes in enumerate(pages):
                from tools.ocr import ocr_image as ocr_fn
                ocr_text = ocr_fn(img_bytes)
                if ocr_text.strip():
                    parts.append(f'[第{i + 1}页]\n{ocr_text}')
            result_text = '\n\n'.join(parts) if parts else ''
            return {"success": True, "text": result_text, "filename": filename, "method": "pdf_ocr", "route": "image"}

        # ── 图片 ──
        if ext in IMAGE_EXTENSIONS:
            use_vision = False
            if mode == 'vision':
                use_vision = True
            elif mode == 'auto':
                from tools.vision import classify_image_complexity
                complexity = classify_image_complexity(content)
                use_vision = (complexity == 'complex')
            else:
                use_vision = False

            if use_vision:
                from tools.vision import VisionAnalyzer
                analyzer = VisionAnalyzer()
                prompt = (
                    "请分析这张图片，提取所有与诈骗案件相关的信息，包括但不限于：\n"
                    "1. 涉及的人员信息（姓名、电话、账号等）\n"
                    "2. 资金信息（金额、转账记录、银行等）\n"
                    "3. 诈骗类型和手法描述\n"
                    "4. 时间线信息\n"
                    "5. 其他所有可见的文字内容和关键信息\n\n"
                    "请以结构化文字形式完整输出所有可识别的信息。"
                )
                result = analyzer.analyze(content, prompt, format=ext.lstrip('.'))
                raw_text = result.get("text", "")
                cleaned_text = _clean_chat_text(raw_text) if clean == 'auto' and raw_text else raw_text
                return {
                    "success": True,
                    "text": cleaned_text,
                    "raw_text": raw_text if cleaned_text != raw_text else "",
                    "filename": filename,
                    "method": result.get("model", "vision"),
                    "route": "vision",
                    "cleaned": cleaned_text != raw_text
                }
            else:
                from tools.ocr import ocr_image as ocr_fn
                text = ocr_fn(content)
                cleaned_text = _clean_chat_text(text) if clean == 'auto' and text else text
                return {
                    "success": True,
                    "text": cleaned_text,
                    "raw_text": text if cleaned_text != text else "",
                    "filename": filename,
                    "method": "ocr",
                    "route": "image",
                    "cleaned": cleaned_text != text
                }
    
        return JSONResponse(status_code=400, content={
            "success": False, "error": f"不支持的文件格式: {ext}"
        })
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


# ──────────────────────────────────────────────
#  端点4: 视觉分析（纯多模态，不降级）
# ──────────────────────────────────────────────
@router.post('/vision-analyze')
async def api_vision_analyze(
    file: UploadFile = File(...),
    prompt: str = Query('请详细描述这张图片的内容', description='分析提示词'),
    clean: str = Query('auto', description='auto: 自动AI清洗结果, off: 不清洗'),
    current_user: dict = Depends(get_current_user)
):
    """完全使用多模态大模型分析图片，适合复杂截图/表格"""
    try:
        filename = file.filename or ''
        ext = os.path.splitext(filename)[1].lower()
        if ext not in IMAGE_EXTENSIONS:
            return JSONResponse(status_code=400, content={
                "success": False, "error": f"视觉分析仅支持图片格式, 不支持: {ext}"
            })
        from tools.vision import VisionAnalyzer
        content = await file.read()
        analyzer = VisionAnalyzer()
        result = analyzer.analyze(content, prompt, format=ext.lstrip('.'))
        raw_text = result.get("text", "")
        cleaned_text = _clean_chat_text(raw_text) if clean == 'auto' and raw_text else raw_text
        return {
            "success": True,
            "text": cleaned_text,
            "raw_text": raw_text if cleaned_text != raw_text else "",
            "filename": filename,
            "model": result.get("model", "unknown"),
            "note": result.get("note", ""),
            "cleaned": cleaned_text != raw_text
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


# ──────────────────────────────────────────────
#  端点5: 资金流水批量导入（真实材料接入 Phase4）
# ──────────────────────────────────────────────
@router.post('/import-fund-flow')
async def api_import_fund_flow(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """导入银行/AMLSim 资金流水 CSV/Excel → accounts_tx，供 /agent-analyze 资金链/回流闭环消费。

    返回 accounts_tx 列表 + 统计；best-effort 落库 ImportedFundFlow 表（合规留痕），失败不阻塞导入。
    """
    from gnn.adapters.fund_flow_io import parse_fund_flow_file
    from database.models import ImportedFundFlow
    from database import db
    try:
        filename = file.filename or 'flow.csv'
        content = await file.read()
        accounts_tx, stats = parse_fund_flow_file(filename, content)
        if not accounts_tx:
            return JSONResponse(status_code=400, content={
                "success": False,
                "error": "未解析到任何资金流转记录（请检查列名是否含发送/接收账户，如 "
                         "SENDER_ACCOUNT_ID/RECEIVER_ACCOUNT_ID 或 付款账号/收款账号）"
            })
        # best-effort 落库（合规留痕）；失败仅告警，不阻塞研判
        try:
            ImportedFundFlow.__table__.create(bind=db.engine, checkfirst=True)
            for tx in accounts_tx:
                db.session.add(ImportedFundFlow(
                    operator=current_user.get('username', ''),
                    source_file=filename,
                    from_account=str(tx.get('from_account')),
                    to_account=str(tx.get('to_account')),
                    amount=float(tx.get('amount') or 0),
                    tx_timestamp=str(tx.get('timestamp') or ''),
                    raw=tx,
                ))
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            from tools.response import logger
            logger.warning(f"[FundFlow] 落库失败（已跳过，不影响本次研判）: {e}")
        object_key, object_url = _store_raw_to_minio(filename, content)
        return {
            "success": True,
            "accounts_tx": accounts_tx,
            "stats": stats,
            "object_key": object_key,
            "object_url": object_url,
            "note": "已将资金流水解析为 accounts_tx；可随 /agent-analyze 的 accounts_tx 字段提交，参与资金链/回流闭环研判",
        }
    except Exception as e:
        return JSONResponse(status_code=400, content={"success": False, "error": str(e)})


# ──────────────────────────────────────────────
#  端点6: 案卷材料 → 结构化 cases（B-L6 4.1 真实材料接入）
# ──────────────────────────────────────────────
@router.post('/parse-case')
async def api_parse_case(
    file: UploadFile = File(None),
    text: str = Query(None, description='直接传案卷文本（与 file 二选一）'),
    current_user: dict = Depends(get_current_user),
):
    """B-L6 4.1：案卷材料（docx/pdf/图片/文本）→ 结构化 cases。

    返回主链路可消费的 cases 列表（含 extracted_entities 供聚类关联）。
    默认零出域：仅本地正则 + 文档标注行解析，不调 LLM。
    诚实边界：输出 source 固定为 "extracted_from_document"，脱敏账户如实保留、不补全。
    """
    try:
        content = None
        filename = ""
        raw_text = None
        if file is not None:
            filename = file.filename or ''
            ext = os.path.splitext(filename)[1].lower()
            content = await file.read()
            if ext == '.docx':
                raw_text = extract_docx_in_order(content)
            elif ext == '.pdf':
                raw_text = _extract_pdf_content(content)
            elif ext in IMAGE_EXTENSIONS:
                from tools.ocr import ocr_image
                ocr_txt = ocr_image(content)
                raw_text = _clean_chat_text(ocr_txt) if ocr_txt else ocr_txt
            elif ext in TEXT_EXTENSIONS:
                raw_text = content.decode('utf-8', errors='ignore')
            else:
                return JSONResponse(status_code=400, content={
                    "success": False, "error": f"不支持的格式: {ext}（支持 docx/pdf/图片/csv/txt）"
                })
            source = f"file:{ext.lstrip('.')}"
        elif text:
            raw_text = text
            filename = "inline_text"
            source = "text"
        else:
            return JSONResponse(status_code=400, content={
                "success": False, "error": "请提供 file（上传材料）或 text（直接文本）之一"
            })

        if not raw_text or not raw_text.strip():
            return JSONResponse(status_code=400, content={
                "success": False,
                "error": "未能从材料中提取到文本（图片 OCR 可能失败，请重试或改用文本/PDF）"
            })

        from gnn.adapters.case_document_io import parse_case_document, extract_docx_in_order
        cases = parse_case_document(raw_text)

        object_key, object_url = (None, None)
        if content is not None:
            object_key, object_url = _store_raw_to_minio(filename, content)

        stats = {
            "n_cases": len(cases),
            "scam_types": _dedupe([c["scam_type"] for c in cases if c["scam_type"] != "未知"]),
            "total_accounts": sum(len(c["extracted_entities"].get("bank_accounts", [])) for c in cases),
            "total_phones": sum(len(c["extracted_entities"].get("phone_numbers", [])) for c in cases),
        }
        return {
            "success": True,
            "cases": cases,
            "stats": stats,
            "source": source,
            "filename": filename,
            "object_key": object_key,
            "object_url": object_url,
            "note": "已将案卷材料结构化为主链路 cases；可随 /agent-analyze 的 cases 字段提交。脱敏账户如实保留，未补全。",
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@router.get('/object/{key:path}')
async def api_get_object(key: str, current_user: dict = Depends(get_current_user)):
    """A4.1 (#C13): 经后端代理下载 minio 对象（数据不出域，前端不直连 minio）。"""
    from core.object_store import get_object, is_enabled
    from fastapi.responses import Response
    if not is_enabled():
        return JSONResponse(status_code=404, content={"success": False, "error": "对象存储不可用"})
    data = get_object(key)
    if data is None:
        return JSONResponse(status_code=404, content={"success": False, "error": "对象不存在"})
    fname = key.split('/')[-1]
    return Response(content=data, media_type="application/octet-stream",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})