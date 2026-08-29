import os
from datetime import datetime
from . import db
from .crud import get_case_by_id, get_gang_by_id
from .models import Case, Gang, GangCaseRelation, EvidenceItem
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, HRFlowable)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'reports')

# 中文公文用宋体。reportlab 自带 STSong-Light CID 字体，无需外部字体文件，
# 避免服务器缺字体导致中文变方框（此前 PDF 正文用 Helvetica，中文全是豆腐块，很丑）。
_CN_FONT = 'Helvetica'
try:
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    _CN_FONT = 'STSong-Light'
except Exception:
    _CN_FONT = 'Helvetica'

# 机关红（GB/T 公文红色标题），与密级黑
_RED = colors.HexColor('#c00000')
_BLACK = colors.HexColor('#1a1a1a')


def _ensure_reports_dir():
    os.makedirs(REPORTS_DIR, exist_ok=True)


def _get_styles():
    styles = getSampleStyleSheet()
    # 红色机关名（红头）
    styles.add(ParagraphStyle(
        'RedOrg',
        parent=styles['Title'],
        fontName=_CN_FONT,
        fontSize=26,
        leading=34,
        alignment=TA_CENTER,
        textColor=_RED,
        spaceAfter=6,
        spaceBefore=4
    ))
    # 文书标题（黑色大字）
    styles.add(ParagraphStyle(
        'ChineseTitle',
        parent=styles['Title'],
        fontName=_CN_FONT,
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        textColor=_BLACK,
        spaceAfter=10
    ))
    styles.add(ParagraphStyle(
        'ChineseHeading',
        parent=styles['Heading2'],
        fontName=_CN_FONT,
        fontSize=14,
        leading=20,
        textColor=_BLACK,
        spaceBefore=12,
        spaceAfter=6
    ))
    styles.add(ParagraphStyle(
        'ChineseBody',
        parent=styles['Normal'],
        fontName=_CN_FONT,
        fontSize=10.5,
        leading=18,
        textColor=_BLACK,
        alignment=TA_JUSTIFY,
        firstLineIndent=21,
        spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        'ChineseMeta',
        parent=styles['Normal'],
        fontName=_CN_FONT,
        fontSize=10,
        leading=16,
        textColor=_BLACK,
        alignment=TA_CENTER,
        spaceAfter=2
    ))
    styles.add(ParagraphStyle(
        'SignRight',
        parent=styles['Normal'],
        fontName=_CN_FONT,
        fontSize=11,
        leading=22,
        textColor=_BLACK,
        alignment=TA_RIGHT,
        spaceBefore=18
    ))
    styles.add(ParagraphStyle(
        'ChineseSmall',
        parent=styles['Normal'],
        fontName=_CN_FONT,
        fontSize=8,
        leading=12,
        textColor=colors.grey
    ))
    return styles


def _report_no(prefix, obj_id):
    """生成公文式文号：FraudLens〔2026〕案研字第 0001 号风格。"""
    year = datetime.now().year
    seq = (int(datetime.now().timestamp()) % 9000) + 1000
    return f'FraudLens〔{year}〕{prefix}字第 {seq} 号'


def _build_red_header(styles, doc_title, report_no, secret='机密', extra_meta=None):
    """红头 + 标题 + 文号/密级行 + 红色分隔线。返回列表 flowables。"""
    els = []
    els.append(Paragraph('FraudLens 反诈智能研判系统', styles['RedOrg']))
    els.append(Spacer(1, 2 * mm))
    els.append(Paragraph(doc_title, styles['ChineseTitle']))
    els.append(Spacer(1, 3 * mm))
    sep = '　　'  # 全角空格分隔，reportlab 不支持 &nbsp;
    meta_line = (f'{report_no}{sep}密级：<font color="#c00000"><b>{secret}</b></font>{sep}'
                 f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    if extra_meta:
        meta_line = extra_meta + sep + meta_line
    els.append(Paragraph(meta_line, styles['ChineseMeta']))
    els.append(Spacer(1, 3 * mm))
    # 红色武断线（红头文件标志性红线）
    els.append(HRFlowable(width='100%', thickness=2, color=_RED, spaceBefore=0, spaceAfter=8))
    return els


def generate_case_report(case_id):
    _ensure_reports_dir()
    case = get_case_by_id(case_id)
    if not case:
        raise ValueError(f'Case {case_id} not found')

    file_name = f'report_case_{case_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    file_path = os.path.join(REPORTS_DIR, file_name)

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        topMargin=20*mm,
        bottomMargin=20*mm,
        leftMargin=15*mm,
        rightMargin=15*mm
    )
    styles = _get_styles()
    elements = []

    # 红头 + 文书标题 + 文号/密级 + 红线（仿公安办案文书版式）
    elements.extend(_build_red_header(
        styles,
        '反诈智能研判报告',
        _report_no('案研', case_id),
        secret='机密',
        extra_meta=f'案件编号：{case_id}'
    ))

    elements.append(Paragraph('一、案件基本信息', styles['ChineseHeading']))
    info_data = [
        ['案件标题', case.get('title', '')],
        ['诈骗类型', case.get('scam_type', '')],
        ['风险等级', case.get('risk_label', '')],
        ['涉案金额', case.get('amount', '')],
        ['案件来源', case.get('source', '')],
        ['诈骗手段', case.get('description', '')[:200] if case.get('description') else ''],
    ]
    info_table = Table(info_data, colWidths=[35*mm, 105*mm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), _CN_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 6*mm))

    elements.append(Paragraph('二、受害人信息', styles['ChineseHeading']))
    victim_data = [
        ['姓名', case.get('victim', '')],
        ['性别', case.get('victim_gender', '')],
        ['年龄', case.get('victim_age', '')],
        ['电话', case.get('victim_phone', '')],
        ['职业', case.get('victim_job', '')],
        ['地址', case.get('victim_address', '')],
    ]
    victim_table = Table(victim_data, colWidths=[35*mm, 105*mm])
    victim_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), _CN_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(victim_table)
    elements.append(Spacer(1, 6*mm))

    entities = case.get('extracted_entities', {})
    if entities:
        elements.append(Paragraph('三、提取实体信息', styles['ChineseHeading']))
        entity_rows = []
        for key, values in entities.items():
            if isinstance(values, list) and values:
                entity_rows.append([key, ', '.join(str(v) for v in values[:5])])
        if entity_rows:
            entity_table = Table(entity_rows, colWidths=[35*mm, 105*mm])
            entity_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), _CN_FONT),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            elements.append(entity_table)
            elements.append(Spacer(1, 6*mm))

    keywords = case.get('keywords', [])
    if keywords:
        elements.append(Paragraph('四、关键词分析', styles['ChineseHeading']))
        kw_text = '、'.join(keywords)
        elements.append(Paragraph(kw_text, styles['ChineseBody']))
        elements.append(Spacer(1, 4*mm))

    evidence_list = EvidenceItem.query.filter_by(case_id=case_id).all()
    if evidence_list:
        elements.append(Paragraph('五、证据材料', styles['ChineseHeading']))
        evi_data = [['编号', '类型', '内容', '状态']]
        for i, ev in enumerate(evidence_list, 1):
            evi_data.append([str(i), ev.type, (ev.content or '')[:60], ev.status])
        evi_table = Table(evi_data, colWidths=[12*mm, 30*mm, 68*mm, 25*mm])
        evi_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), _CN_FONT),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(evi_table)
        elements.append(Spacer(1, 6*mm))

    ai_report = case.get('ai_report', '')
    if ai_report:
        elements.append(Paragraph('六、智能分析报告', styles['ChineseHeading']))
        # 保留 AI 报告换行：转成 <br/> 再用 Paragraph 渲染
        safe = (ai_report[:3000].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                .replace('\n', '<br/>'))
        elements.append(Paragraph(safe, styles['ChineseBody']))

    # 承办人签章位（公文落款，右对齐）
    elements.append(Paragraph('研判民警（承办）：__________________　复核：__________________', styles['SignRight']))
    elements.append(Paragraph(datetime.now().strftime('%Y年%m月%d日'), styles['SignRight']))
    elements.append(Spacer(1, 8*mm))
    elements.append(HRFlowable(width='100%', thickness=0.6, color=colors.grey))
    elements.append(Paragraph('本报告由 FraudLens 反诈智能研判系统自动生成，数据来源于本单位授权系统，仅供内部研判参考，严禁外传。', styles['ChineseSmall']))

    doc.build(elements)
    return file_path


def generate_gang_report(gang_id):
    _ensure_reports_dir()
    gang = get_gang_by_id(gang_id)
    if not gang:
        raise ValueError(f'Gang {gang_id} not found')

    file_name = f'report_gang_{gang_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    file_path = os.path.join(REPORTS_DIR, file_name)

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        topMargin=20*mm,
        bottomMargin=20*mm,
        leftMargin=15*mm,
        rightMargin=15*mm
    )
    styles = _get_styles()
    elements = []

    # 红头 + 文书标题 + 文号/密级 + 红线
    elements.extend(_build_red_header(
        styles,
        '涉诈团伙研判报告',
        _report_no('团研', gang_id),
        secret='机密',
        extra_meta=f'团伙编号：{gang_id}'
    ))

    elements.append(Paragraph('一、团伙基本信息', styles['ChineseHeading']))
    info_data = [
        ['团伙名称', gang.get('gang_name', '')],
        ['风险等级', gang.get('risk_label', '')],
        ['综合评分', str(gang.get('comprehensive_score', 0))],
        ['可信度', str(gang.get('confidence', 0))],
        ['预估成员数', gang.get('member_count_estimate', '')],
        ['技术等级', gang.get('tech_level', '')],
        ['剧本类型', gang.get('script_type', '')],
        ['关联案件数', str(gang.get('total_cases', 0))],
        ['总涉案金额', gang.get('total_amount_involved', '')],
    ]
    info_table = Table(info_data, colWidths=[35*mm, 105*mm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), _CN_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 6*mm))

    related = gang.get('related_cases', [])
    if related:
        elements.append(Paragraph('二、关联案件', styles['ChineseHeading']))
        case_data = [['编号', '案件ID', '受害人', '涉案金额', '风险等级']]
        for i, rc in enumerate(related, 1):
            case_data.append([
                str(i),
                rc.get('case_id', ''),
                rc.get('victim', ''),
                rc.get('amount', ''),
                rc.get('risk_level', ''),
            ])
        case_table = Table(case_data, colWidths=[10*mm, 30*mm, 35*mm, 35*mm, 30*mm])
        case_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), _CN_FONT),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(case_table)
        elements.append(Spacer(1, 6*mm))

    description = gang.get('description', '')
    if description:
        elements.append(Paragraph('三、团伙描述', styles['ChineseHeading']))
        elements.append(Paragraph(description[:1000], styles['ChineseBody']))
        elements.append(Spacer(1, 6*mm))

    modus_operandi = gang.get('modus_operandi', '')
    if modus_operandi:
        elements.append(Paragraph('四、作案手法', styles['ChineseHeading']))
        elements.append(Paragraph(modus_operandi[:1000], styles['ChineseBody']))
        elements.append(Spacer(1, 6*mm))

    fingerprint = gang.get('fingerprint', [])
    if fingerprint:
        elements.append(Paragraph('五、团伙特征指纹', styles['ChineseHeading']))
        for fp in fingerprint:
            if isinstance(fp, str):
                elements.append(Paragraph(f'• {fp}', styles['ChineseBody']))
            elif isinstance(fp, dict):
                elements.append(Paragraph(f'• {fp.get("name", "")}: {fp.get("value", "")}', styles['ChineseBody']))

    prevention_advice = gang.get('prevention_advice', '')
    if prevention_advice:
        elements.append(Paragraph('六、防范建议', styles['ChineseHeading']))
        if isinstance(prevention_advice, list):
            for item in prevention_advice:
                elements.append(Paragraph(f'• {item}', styles['ChineseBody']))
        else:
            elements.append(Paragraph(str(prevention_advice)[:1000], styles['ChineseBody']))

    risk_assessment = gang.get('risk_assessment', {})
    if risk_assessment:
        elements.append(Paragraph('七、风险评估', styles['ChineseHeading']))
        if isinstance(risk_assessment, dict):
            for key, value in risk_assessment.items():
                elements.append(Paragraph(f'• {key}: {value}', styles['ChineseBody']))
        else:
            elements.append(Paragraph(str(risk_assessment)[:1000], styles['ChineseBody']))

    steps = gang.get('steps', [])
    if steps:
        chain = ' → '.join(s if isinstance(s, str) else (s.get('title') or s.get('name') or s.get('step') or '')
                           for s in steps)
        if chain.strip(' →'):
            elements.append(Paragraph('八、作案流程链', styles['ChineseHeading']))
            elements.append(Paragraph(chain, styles['ChineseBody']))

    # 承办人签章位（公文落款）
    elements.append(Paragraph('研判民警（承办）：__________________　复核：__________________', styles['SignRight']))
    elements.append(Paragraph(datetime.now().strftime('%Y年%m月%d日'), styles['SignRight']))
    elements.append(Spacer(1, 8*mm))
    elements.append(HRFlowable(width='100%', thickness=0.6, color=colors.grey))
    elements.append(Paragraph('本报告由 FraudLens 反诈智能研判系统自动生成，数据来源于本单位授权系统，仅供内部研判参考，严禁外传。', styles['ChineseSmall']))

    doc.build(elements)
    return file_path


def export_case_docx(case_id):
    _ensure_reports_dir()
    case = get_case_by_id(case_id)
    if not case:
        raise ValueError(f'Case {case_id} not found')

    file_name = f'case_{case_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    file_path = os.path.join(REPORTS_DIR, file_name)

    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = document.add_heading('反诈智能研判报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'报告编号: RPT-{case_id}    生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    document.add_heading('一、案件基本信息', level=1)
    info_table = document.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('案件标题', case.get('title', '') or ''),
        ('诈骗类型', case.get('scam_type', '') or ''),
        ('风险等级', case.get('risk_label', '') or ''),
        ('涉案金额', case.get('amount', '') or ''),
        ('案件来源', case.get('source', '') or ''),
        ('案件状态', case.get('status', '') or ''),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label or ''
        info_table.rows[i].cells[1].text = value or ''

    document.add_heading('二、受害人信息', level=1)
    victim_table = document.add_table(rows=6, cols=2)
    victim_table.style = 'Table Grid'
    victim_data = [
        ('姓名', case.get('victim', '') or ''),
        ('性别', case.get('victim_gender', '') or ''),
        ('年龄', case.get('victim_age', '') or ''),
        ('电话', case.get('victim_phone', '') or ''),
        ('职业', case.get('victim_job', '') or ''),
        ('地址', case.get('victim_address', '') or ''),
    ]
    for i, (label, value) in enumerate(victim_data):
        victim_table.rows[i].cells[0].text = label or ''
        victim_table.rows[i].cells[1].text = value or ''

    entities = case.get('extracted_entities', {})
    if entities:
        document.add_heading('三、提取实体信息', level=1)
        valid_entities = {k: v for k, v in entities.items() if isinstance(v, list) and v}
        if valid_entities:
            entity_table = document.add_table(rows=len(valid_entities), cols=2)
            entity_table.style = 'Table Grid'
            for i, (key, values) in enumerate(valid_entities.items()):
                entity_table.rows[i].cells[0].text = key
                entity_table.rows[i].cells[1].text = ', '.join(str(v) for v in values[:5])

    keywords = case.get('keywords', [])
    if keywords:
        document.add_heading('四、关键词分析', level=1)
        document.add_paragraph('、'.join(keywords))

    ai_report = case.get('ai_report', '')
    if ai_report:
        document.add_heading('五、智能分析报告', level=1)
        document.add_paragraph(ai_report[:2000])

    document.save(file_path)
    return file_path