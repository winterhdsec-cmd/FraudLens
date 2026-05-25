"""FraudLens 支撑材料 Word - 侧重成品展示"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'FraudLens_支撑材料_v2.docx')
doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def add_h1(text):
    _add_empty()
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.name = '黑体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.line_spacing = 1.5
    _add_empty()

def add_h2(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.name = '楷体_GB2312'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(6)

def add_body(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name = '宋体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(2)

def add_figure(fig_num, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'【此处插入图（{fig_num}）：{caption}】')
    r.font.name = '宋体'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'图（{fig_num}）：{caption}')
    r2.bold = True; r2.font.name = '宋体'; r2.font.size = Pt(11)
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(8)

def add_empty():
    p = doc.add_paragraph(); r = p.add_run(''); r.font.size = Pt(6)
    p.paragraph_format.line_spacing = 1.0

_add_empty = add_empty

# ══════════════════ 封面 ══════════════════
for _ in range(4): add_empty()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('支撑材料'); r.bold = True; r.font.name = '黑体'; r.font.size = Pt(26)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
add_empty(); add_empty()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('多智能体协同反欺诈智能研判系统')
r.bold = True; r.font.name = '黑体'; r.font.size = Pt(22)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FraudLens v3.1'); r.font.name = '宋体'; r.font.size = Pt(16)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
r.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
for _ in range(4): add_empty()
for line in ['项目名称：多智能体协同反欺诈智能研判系统','项目编号：_____________','所属类别：计算机类','项目级别：_____________']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.name = '宋体'; r.font.size = Pt(14)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.add_page_break()

# ══ 一、系统简介 ══
add_h1('一、系统简介')
add_body('FraudLens（多智能体协同反欺诈智能研判系统）是一款面向公安机关反诈中心的人工智能研判系统。系统集案件管理、智能分析、团伙挖掘、资金追踪、预警监控于一体，支持文本、图片、文档等多种数据录入方式，通过多智能体协同流水线自动完成案件分析研判。系统基于FastAPI + Vue 3构建前后端分离架构，采用BGE语义编码与HDBSCAN聚类算法实现诈骗团伙自动发现，并利用大语言模型生成智能研判报告。')
doc.add_page_break()

# ══ 二、系统功能展示 ══
add_h1('二、系统功能展示')

add_h2('2.1 案件管理与智能分析')
add_body('系统提供完整的案件管理功能，支持手动录入、批量导入、文件智能解析等多种数据接入方式。录入完成后一键启动智能分析，系统通过六阶段多智能体流水线自动完成案件分类、风险分级、实体提取、团伙发现等任务，分析过程通过WebSocket实时推送进度。案件详情页包含概览信息、行为特征雷达图、资金流向、AI研判报告等多个分析维度，全方位呈现案件特征。')
add_figure('一', '系统总览看板页面截图')
add_figure('二', '案件智能分析结果页面截图')

add_h2('2.2 团伙画像深度分析')
add_body('团伙画像是系统的核心功能之一。系统基于语义编码和密度聚类算法自动发现犯罪团伙，并生成完整的团伙画像信息。画像包含团伙基本信息（名称、风险等级、成员规模、涉案金额）、特征雷达图（从诈骗话术成熟度、资金分散程度、技术手段先进性、跨区域作案特征、反侦察能力、受害者画像精准度六个维度进行量化评估）、作案手法描述以及预防建议。用户可点击团伙卡片进入深度分析页面，查看团伙关联的所有案件和资金流向。')
add_figure('三', '团伙画像深度分析与特征雷达图截图')

add_h2('2.3 资金流向追踪')
add_body('资金流向追踪模块以可视化图形展示诈骗资金的完整流转路径。系统自动将资金流转分为多个层级，从受害人向嫌疑人账户首次转账开始，经过多级账户中转、地下钱庄、第三方支付、虚拟货币兑换等复杂洗钱路径，并利用AI自动生成每笔流转的文字描述。提供资金统计概览面板，展示涉案总金额、涉及账户数、最大层级深度、境外资金比例等关键指标。')
add_figure('四', '资金流向追踪可视化界面截图')

add_h2('2.4 关联图谱与预警监控')
add_body('关联图谱模块基于vis-network力导向图，展示案件、团伙、人员之间的复杂关联关系，节点颜色和大小根据风险等级动态编码，支持交互式拖拽与筛选。预警监控模块利用Redis实时缓存和实体匹配算法，新案件录入时自动比对手机号、银行卡号、APP名称、IP地址等关键实体，发现匹配项即生成含置信度分数的预警记录。')
add_figure('五', '关联图谱与预警监控界面截图')

add_h2('2.5 报告导出与会话管理')
add_body('系统支持将分析结果导出为PDF和Word格式标准报告，涵盖案件基本信息、资金流向、AI分析结论、团伙关联、侦查建议等完整内容。同时提供分析会话管理功能，跟踪每次分析的进度和结果，支持历史会话回溯。')
add_figure('六', '报告导出与系统管理界面截图')

doc.add_page_break()

# ══ 三、使用方法 ══
add_h1('三、使用方法')

add_h2('3.1 系统登录与案件录入')
add_body('系统提供Web端访问方式，用户在浏览器中打开系统地址后进入登录页面，输入账号密码即可进到系统主界面。在主界面左侧导航栏点击"案件管理"进入案件列表页，可通过"手动录入"填写案件信息并保存，或通过"批量导入"上传Excel文件实现多案件批量录入。系统同时支持图片和文档解析功能，可自动提取文件中的案件要素信息。')
add_figure('七', '系统登录页面与案件录入界面截图')

add_h2('3.2 智能分析与团伙研判')
add_body('录入案件后，在案件列表页选中目标案件点击"开始分析"，系统即进入多智能体协同分析流水线。分析过程通过WebSocket实时推送进度，用户可在案件详情页的"分析进度"标签页中实时查看各阶段完成状态。分析完成后自动生成案件风险评级、实体提取结果和AI研判报告。在"团伙画像"模块中，系统将语义相似的案件自动归并为犯罪团伙并展示团伙详情，用户可点击团伙名称进入深度分析页面查看团伙特征雷达图、关联案件列表和资金流向全貌。')
add_figure('八', '智能分析进度展示与团伙画像页面截图')

add_h2('3.3 资金追踪与报告导出')
add_body('在案件详情页点击"资金流向"标签页，系统以可视化图形式展示资金流转路径全貌。使用鼠标滚轮可缩放查看，拖拽节点可调整布局。在"报告导出"功能中，用户可选择导出为PDF或Word格式，系统将自动生成包含案件基本信息、资金流向、AI分析结论、侦查建议的完整分析报告。')
add_body('（1）案件录入：通过手动录入、批量导入或文件解析方式添加案件数据。（2）启动分析：选中案件一键启动多智能体协同分析流水线。（3）查看结果：在案件详情页查看分析结果，在团伙画像页查看团伙关联。（4）导出报告：将分析结果导出为PDF或Word格式文件。')

doc.add_page_break()

# ══ 四、系统技术架构 ══
add_h1('四、系统技术架构')
add_body('系统基于FastAPI + Vue 3前后端分离架构。后端采用Python 3.11，集成DeepSeek Chat大语言模型、BGE语义编码模型、EasyOCR引擎、HDBSCAN聚类算法等AI能力；前端使用Vue 3 + Element Plus + ECharts构建可视化交互界面。数据层采用MySQL 8.0 + Redis 7，通过Docker容器化部署。')
add_body('系统的核心技术亮点包括：（1）多智能体协同流水线，由ChiefAgent协调五个专业Agent完成从数据预处理到画像增强的全流程分析；（2）BGE语义编码与UMAP降维HDBSCAN聚类，自动将语义相似案件归并为犯罪团伙；（3）增量聚类机制，支持新案件持续归入已有团伙的动态更新；（4）AI自动资金流向标注与实时实体匹配预警。')
add_figure('九', '系统总体技术架构图')
add_figure('十', '多智能体协同与聚类算法流程图')

doc.add_page_break()

# ══ 五、成品图片 ══
add_h1('五、成品图片')
add_body('以下是系统实际运行界面展示。')

add_figure('十一', '系统界面展示——总览看板与案件详情')
add_body('总览看板展示整体态势统计，案件详情页展示案件概览信息和行为特征雷达图。')

add_figure('十二', '系统界面展示——团伙画像与资金流向')
add_body('团伙深度分析页面包含特征雷达图及其关联案件，资金流向页面展示多层级资金流转链路。')

doc.save(OUTPUT)
print(f'文档已生成: {OUTPUT}')
