# -*- coding: utf-8 -*-
"""CPEC2026 匿名版论文 Word 生成脚本
- 按《附件2：论文模版》格式：五号宋体正文、黑体标题、200字摘要、三线表、图题居中
- 匿名版：不署名、不写基金/作者简介
- 目标 4-5 页 A4
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'E:/FraudLens/paper/论文产出/CPEC2026教学案例稿/面向反诈团伙研判的AI赋能实验教学案例设计_匿名版.docx'

doc = Document()

# ===== 页面设置：A4，标准页边距 =====
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(3.0)

# ===== 默认字体：宋体五号 =====
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15
pf.space_after = Pt(0)
pf.space_before = Pt(0)

def set_font(run, east='宋体', ascii_f='Times New Roman', size=10.5, bold=False):
    run.font.name = ascii_f
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east)

def para(text='', align=None, size=10.5, bold=False, east='宋体',
         indent_first=None, space_after=0, space_before=0, line=1.3):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if indent_first:
        pf.first_line_indent = Pt(indent_first)
    if text:
        run = p.add_run(text)
        set_font(run, east=east, size=size, bold=bold)
    return p

def heading1(num, text):
    p = para(f'{num} {text}', align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True, east='黑体',
             space_before=8, space_after=4, line=1.3)
    return p

def heading2(num, text):
    p = para(f'{num} {text}', align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, bold=True, east='黑体',
             space_before=4, space_after=2, line=1.3)
    return p

def body(text, indent=True):
    p = para(text, size=10.5, indent_first=21 if indent else None, line=1.15)
    return p

def caption_zh(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, bold=True, east='黑体',
                space_before=2, space_after=2, line=1.1)

def caption_en(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, east='Times New Roman',
                space_before=0, space_after=4, line=1.1)

def make_3line_table(headers, rows, col_widths=None, font_size=9):
    """三线表：顶线/栏目线/底线，用边框实现"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 写表头
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, east='黑体', size=font_size, bold=True)
    # 写数据
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(v))
            set_font(run, east='宋体', size=font_size)
    # 边框：顶部+栏目线+底部
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for edge, sz in [('top', 12 if i == 0 else 0),
                             ('bottom', 12 if i == len(t.rows) - 1 else 6 if i == 0 else 0),
                             ('left', 0), ('right', 0)]:
                el = OxmlElement(f'w:{edge}')
                if sz:
                    el.set(qn('w:val'), 'single')
                    el.set(qn('w:sz'), str(sz))
                    el.set(qn('w:color'), '000000')
                else:
                    el.set(qn('w:val'), 'nil')
                borders.append(el)
            tcPr.append(borders)
    return t

# ============================================================
# 题名区（匿名版：无作者/单位/基金）
# ============================================================
para('中图分类号：TP393.08；TP18；G642　　文献标志码：A',
     align=WD_ALIGN_PARAGRAPH.LEFT, size=9, space_after=6)

para('面向反诈团伙研判的 AI 赋能实验教学案例设计',
     align=WD_ALIGN_PARAGRAPH.CENTER, size=15, bold=True, east='黑体', space_after=6)

# ---- 中文摘要（≤200字）----
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.15
r1 = p.add_run('摘　要：')
set_font(r1, east='黑体', bold=True)
abstract = ('针对公安院校反诈实战教学存在学生缺乏真实系统操作机会、教师缺少可复现实训材料的问题，'
            '以可运行的反诈团伙研判原型系统 FraudLens 为载体，设计 AI 赋能实验教学案例。'
            '案例将“案件级串并案＋账户级扩线”两级研判闭环改写为实训内容，'
            '按“研—训—评”模式组织案情分析、工具辅助串并案、冻卡决策、边界反思四个递进实验环节，'
            '配套带权重的评价量表与思政映射，并以可一键复现的能力基线与诚实边界认知支撑教学。'
            '实训以研判决策而非模型分析为核心，培养学生在 AI 辅助下的研判决策、人机协同与工程可信伦理能力，'
            '为公安实战人才培养提供可复用的领域级实训素材。')
r2 = p.add_run(abstract)
set_font(r2, east='宋体')

para('关键词：反诈教学；图神经网络；实验教学案例；研判决策；人机协同',
     size=10.5, space_after=6)

# ---- 英文题名 / 摘要 ----
para('Design of an AI-Enabled Practical Teaching Case for Anti-Fraud Gang Investigation',
     align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, east='Times New Roman', space_after=4)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.2
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = p.add_run('Abstract: ')
set_font(r1, east='Times New Roman', bold=True)
abstract_en = ('To address the lack of hands-on operation of real investigation systems for students and '
               'the absence of reproducible teaching materials for instructors in anti-fraud practical teaching '
               'at police colleges, this paper designs an AI-enabled practical teaching case based on a running '
               'anti-fraud investigation prototype, FraudLens. The case transforms the two-level investigation loop '
               '(case-level merging and account-level expansion) into laboratory content, and organizes four progressive '
               'experiments—case analysis, AI-assisted case linking, freeze-order decision, and capability-boundary '
               'reflection—under a research–training–evaluation model, supported by one-click reproducible baselines '
               'and honest boundary cognition. Centered on investigative decision-making rather than model analysis, '
               'the case cultivates students\u2019 AI-assisted decision-making, human\u2013AI collaboration, and '
               'engineering trustworthy ethics, providing a deployable domain-level training resource for police talent cultivation.')
r2 = p.add_run(abstract_en)
set_font(r2, east='Times New Roman', size=10)

para('Key words: anti-fraud; graph neural network; practical teaching case; investigative decision-making; human–AI collaboration',
     size=10, space_after=8)

# ============================================================
# 1 引言
# ============================================================
heading1('1', '引言')
body('新工科建设要求计算机类人才具备 AI 赋能的实践能力和工程素养，把真实领域问题改写成可复现的实践教学案例，是提升学生实战能力的一条可行路径。现有计算机实践教学案例大多围绕通用编程和单门课程[1-2]，面向涉网犯罪智能研判的领域级可复现案例仍属稀缺[3]。在公安院校这一问题更为突出：反诈实战课上，学生长期只能听案例、看演示，很难亲手操作真实的研判系统，教师也缺少既适合课堂又能让学生复现的现成材料。')
body('本文以可运行的反诈团伙研判原型系统 FraudLens 为载体，把“案件级串并案＋账户级扩线”两级研判闭环与反思编排整合为实训内容，构建一份可复现的“研—训—评”一体化教学案例。主要工作如下：（1）提供可落地的 AI 辅助研判决策实训载体，其两级研判系统由 LangGraph[4] 编排“规划→预处理→分析→聚类→反思”工作流，技术栈覆盖图建模、异构注意力、系统编排与工程可信；（2）提出“研—训—评”一体化实训模式，按“案情分析→工具辅助串并案→冻卡决策与伦理权衡→边界认知与反思”设计四环节递进实验链，以研判决策为核心活动；（3）提供可复现实验与科学边界认知训练，能力基线（HAN 困难场景 F1=0.9154、反思闭环增益 +0.5125、账户级扩线将团伙 F1 由盲扫约 0.002 提升至约 0.71）全部可一键复现，同时诚实量化反诈 GNN 的增量边界，引导学生建立对模型能力的审慎认知。')
body('文中能力基线数据仅作为实训可复现对象的佐证，不代表真实警务场景的已验证成效；教学成效的课堂实证评估属下一阶段工作。')

# ============================================================
# 2 实训载体：FraudLens 系统概览
# ============================================================
heading1('2', '实训载体：FraudLens 系统概览')
body('FraudLens 是一套可运行的反诈团伙研判原型系统（图1），覆盖数据建模、表示学习、系统编排与工程可信的完整技术链，可作为实训“研”的对象。系统构成两级研判闭环：案件级串并案以异构注意力网络（HAN）[5]在案件异构图上聚合同伙案件；账户级扩线在给定线索锚点后于账户交易子图上还原资金链路，一级输出的关键账户可直接作为二级扩线的锚点，契合真实反诈工作流。')
body('为降低上手门槛，实训以研判决策为核心，不要求学生理解图神经网络内部原理，全部模块均提供预计算结果供学生交互使用。系统的技术构成可概括为四点：异构双通道融合——以 HAN 融合“结构通道（金额加权交易图）”与“文本通道（本地 BGE 语义嵌入）”，支撑案件级团伙聚类；多智能体反思闭环——由 LangGraph[4] 驱动“规划→预处理→分析→聚类→反思”工作流，反思节点在聚类质量不达标时经条件边回连触发真实重算[6]；经验加权置信度门控——以团伙规模、涉案金额、关联账户数、资金回流四因子加权给出冻卡建议置信度；工程可信——实现 RBAC 鉴权、双表审计、多级降级链与“数据不出域”，契合公安院校数据合规要求。')

# 图1
try:
    para('', space_after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'E:/FraudLens/paper/论文产出/CPEC2026教学案例稿/fig_arch.png', width=Cm(11))
    caption_zh('图1  FraudLens 系统总体架构')
except Exception as e:
    print('图1 插入失败:', e)

# ============================================================
# 3 实训可复现性与能力基线
# ============================================================
heading1('3', '实训可复现性与能力基线')
body('本节给出 FraudLens 两级研判能力的可复现性与能力边界的三层验证，作为实训的科学基线：（1）合成数据能力基线与消融——在受控合成案情上确认模型判别水平，定位反思闭环与异构图建模各自的贡献；（2）AMLSim 账户级扩线外部验证——在真实规模的公开反洗钱基准上考察无监督扩线的能力上限；（3）Elliptic 跨数据集边界对照——援引公开基准结果，说明方法对无金额/时序的纯拓扑交易图不迁移。三层验证共同支撑“可复现、有上限、边界诚实”的实训基调，全部数据随系统开放、可一键复现。')

heading2('3.1', '能力基线与消融')
body('表1 报告合成数据（10 随机种子均值）下的能力基线。HAN 在困难（Hard）场景 F1=0.9154，显著高于同构 GraphSAGE（Hard F1=0.3043，因无法建模多类型实体关系发生塌缩），这一对比直接成为实训 Lab2“为何需要 AI 辅助”的教学素材。消融显示，反思闭环缺失使失败场景 F1 由 0.8353 跌至 0.3228（增益 +0.5125），说明反思闭环对该类场景不可或缺；关闭 GNN 退化为 Louvain 社区发现，Hard F1 由 0.9154 降至 0.8616，说明异构图建模对困难场景同样是必要的。')
caption_zh('表1  能力基线与消融（合成数据，10 seeds mean±std）')
make_3line_table(
    ['方法/配置', 'Clean F1', 'Hard F1', '备注'],
    [
        ['HAN（本文）', '1.0000±0.0000', '0.9154±0.1214', '异构注意力融合'],
        ['GraphSAGE', '0.4061±0.2127', '0.3043±0.0000', 'Hard 塌缩（同构局限）'],
        ['Louvain（降级）', '0.8884±0.0911', '0.8616±0.0874', '关闭 GNN 的社区发现'],
        ['w/o 反思闭环', '—', '—', '失败场景 F1 由 0.8353 跌至 0.3228'],
    ]
)
para('注：Clean 场景无跨团伙干扰（cross=0.0），仅验证模型基本学习能力；反思闭环仅作用于失败场景的自动降级重算，故 w/o 反思的 Hard F1 以“—”表示。',
     size=8.5, space_after=6, line=1.15)

heading2('3.2', '账户级扩线外部验证（AMLSim）')
body('为验证账户级扩线（二级闭环）在真实规模公开图上的能力边界，本节在 AMLSim[7] 反洗钱基准（43,614 账户／1,305 个洗钱环，含账户级环真值）上报告无监督扩线结果：给定嫌疑账户锚点，于账户交易图上做 k=1 跳扩线，子图规模约 2,748 节点，全部方法零标签、纯无监督。表2 显示，盲扫（全网无锚点）几乎失效（F1≈0.002），印证“盲扫 vs 扩线”瓶颈；扩线设定下，k-core 核心子图（kmin=2）将团伙识别 F1 提升至约 0.71（召回 0.91），短有向环检测将资金环识别 F1 提升至约 0.62。该结果不宣称真实警务验证——AMLSim 为公开合成基准；其适用边界明确：只适用于含金额/时序的账户交易图，对纯拓扑交易图（Elliptic，无金额时序）不迁移[8]。“无监督有上限、标注能破但不稳”的结论，正是实训 Lab4 的数据素材。')
caption_zh('表2  AMLSim 账户级扩线无监督基线（锚点 k=1 跳，零标签）')
make_3line_table(
    ['任务', '方法', 'Precision', 'Recall', 'F1'],
    [
        ['团伙识别', '盲扫（无锚点）', '—', '—', '≈0.002（失效）'],
        ['团伙识别', 'k-core 核心子图（kmin=2）', '0.58', '0.91', '0.71'],
        ['资金环识别', '短有向环检测（L≤8）', '0.63', '0.61', '0.62'],
    ]
)
para('注：AMLSim 为 IBM 公开合成反洗钱基准；k-core 召回 0.91 但候选占子图约 85%，故 precision 受限，这正是无监督上限的结构性原因。',
     size=8.5, space_after=6, line=1.15)

# ============================================================
# 4 实训教学设计
# ============================================================
heading1('4', '实训教学设计')
body('FraudLens 所覆盖的技术栈——异构图谱建模、双通道语义融合、多智能体反思编排、可解释门控决策——与“新工科”对 AI 赋能实践能力与工程素养的要求能够对应。本节给出将其转化为可复现实训教学案例的设计方案，供“人工智能安全”“图数据挖掘”及相关专业课程参考。')

heading2('4.1', '学情分析与课程衔接')
body('本案例面向公安院校信息技术、网络安全与侦查相关专业本科生（建议大二下至大三上）实施。学生已修 Python 程序设计、数据结构与数据库等课程，具备基本图论概念与编程调试能力，但一般没有图神经网络与多智能体的先验，这两部分不作先修要求，只以黑箱工具形式参与。实训以研判决策为核心活动：图神经网络作为学生交互的黑箱工具，只要求学生理解“系统给出什么结果、可信到什么程度、何时该信、何时不该信”，不要求掌握模型内部机理。全部实验以 Jupyter Notebook 为载体，系统研判结果预烘焙供学生对比分析。')

heading2('4.2', '“研—训—评”一体化模式')
body('本节以“研—训—评”一体化模式组织实训：以真实 FraudLens 系统为“研”的对象，以四环节递进实验链为“训”的载体，以过程＋产物双轨评价及思政/伦理素养观测为“评”的手段，并以可复现实验、诚实边界认知与课程思政作为全程支撑，让育人主线不被技术细节盖过。上述“研—训—评”在教学设计层面已构成完整闭环；实证成效层面的闭环将在 2026 年秋季学期课程实施后闭合。')

heading2('4.3', '四环节递进实验链')
body('实训以四环节递进实验链为主线（表3）：四个 Lab 从案情分析到边界反思逐级递进，每个 Lab 锚定 FraudLens 的真实模块，以研判决策为核心活动。系统研判结果预烘焙为 JSON/NumPy 数组，学生打开 notebook 执行 Run All 即可获得案情数据与系统分析结果，教学时间用于决策、对比与反思。')
caption_zh('表3  四环节递进实验链——课时分配、核心任务与训练产出')
make_3line_table(
    ['Lab', '时长', '核心任务', '对应能力/系统模块', '训练产出'],
    [
        ['1', '4 课时', '分析案情，认识研判流程', '案情分析能力·系统概览', '案情分析报告'],
        ['2', '4 课时', '人工串并案 vs 系统结果', '人机协同·双通道 HAN＋反思闭环', '对比报告'],
        ['3', '4 课时', '门控调参，冻卡决策角色扮演', '决策与伦理·置信度门控', '伦理分析报告'],
        ['4', '4 课时', '系统失败场景与诚实反思', '科学边界认知·外部验证', '边界反思总结'],
    ]
)
para('注：课前预习 2 课时（pip install＋数据集下载，不计入课内学时）；课堂导入与总结内嵌于各 Lab，课内总计 16 学时。',
     size=8.5, space_after=4, line=1.15)

body('Lab1 案情分析与研判流程：学生拿到一份合成案情（含受害者陈述、账户流水、通话记录、涉案金额），以办案人员视角分析涉及的账户、资金流向与关键节点，再由系统展示同一案件的异构图可视化，收尾讨论“如果只用传统逐笔查账方式，你会漏掉什么？”')
body('Lab2 工具辅助串并案：学生先人工判断“哪几个案子可能是同一伙人干的”，再运行系统看聚类结果，对比人与系统各自发现的关联——系统能发现人漏掉的多跳结构关联，人也可能发现系统漏掉的基于办案经验的关联，由此建立“AI 是辅助工具而非替代”的人机协同意识。')
body('Lab3 冻卡决策与伦理权衡：学生调整置信度门控阈值，观察阈值变化如何影响冻卡建议数量，并角色扮演“阈值设高，3 个无辜者账户被冻结；阈值设低，团伙今晚转移资金”，记录决策理由并讨论“算法建议冻，但谁来承担后果”。')
body('Lab4 边界认知与诚实反思：学生观察同构 GNN 基线（GraphSAGE）在困难场景塌缩（F1 由约 0.41 跌至 0.30），而异构 HAN 保持 0.9+；再读 AMLSim 盲扫 F1≈0.002、Elliptic HAN F1≈0.016——真实数据上所有方法都接近失败。引入“扩线”概念：给定锚点后只分析 k 跳邻居，无监督拓扑方法将团伙识别 F1 由盲扫 0.002 提升至约 0.71、资金环识别约 0.62，无需任何标签；但该增益只适用于含金额/时序的账户交易图。教学要点：反诈 GNN 的瓶颈不在算法设计，而在“盲扫 vs 扩线”的任务设定与无监督方法的适用边界。学生撰写 300 字反思：“作为未来的警务技术使用者，你在什么情况下会信任这个系统？什么情况下不会？”')

heading2('4.4', '考核方式设计')
body('考核采用“过程＋产物”双轨：过程维度占 40%（案情分析合理性 15%、人机对比串并案 15%、冻卡决策理由 10%），产物维度占 60%（边界认知反思 30%、伦理分析报告 15%、团队答辩 15%）。评价的核心不是指标最大化，而是学生对“系统能力与局限”的如实把握——既看到 AI 辅助研判的价值，也理解其在真实场景中的局限，这一取向本身即承载科学素养训练。')

heading2('4.5', '课程思政与伦理教育')
body('反诈主题与思政育人天然契合：守护群众财产安全、法治意识、技术向善与总体国家安全观。本实训将思政与伦理元素嵌入各 Lab 而非附加说教：Lab1 强调数据合规与“数据不出域”，Lab2 树立“AI 辅助而非替代民警”的人机协同意识，Lab3 在冻卡决策中开展公民权利保护讨论，Lab4 培育诚实严谨的科学精神。')

heading2('4.6', '教学成效说明与实施预案')
body('本案例设计尚未在正式课程中规模化实施。为使教学成效评估尽量客观，按“预注册”原则预先声明评价指标，待 2026 年秋季学期实施后按同样指标采集数据并如实报告，避免选择性报告偏差。评价体系含六维：案情分析能力（课前/课后问卷）、人机协同意识（Lab2 报告盲评）、决策伦理分析深度（冻卡报告盲评 4 级量表）、诚实边界认知（反思中“合成≠真实”表述准确性）、系统可用性（SUS 量表）与课程推荐意愿（NPS）。')

# ============================================================
# 5 结语
# ============================================================
heading1('5', '结语')
body('本文以可运行的反诈团伙研判原型系统 FraudLens 为载体，设计了面向学生能力成长的 AI 赋能实验教学案例，形成了“研—训—评”一体化的实训模式与四环节递进实验链。区别于聚焦通用编程或单门课程的现有案例，本案例把真实系统转化为可复现的“研”对象，以研判决策为核心组织“训”的链条，以过程＋产物双轨“评”支撑育人成效，并将数据合规、公民权利保护与科学诚信内嵌于实训环节，避免技术训练与价值引领“两张皮”。下一步将在 2026 年秋季学期课程实施中采集六维评价数据，验证案例对 AI 辅助研判决策、人机协同与工程可信伦理能力的培养成效，并据此迭代案例设计。')

# ============================================================
# 参考文献（≥12 篇，正文顺序编码）
# ============================================================
heading1('', '参考文献')
refs = [
    '[1] 李清勇，等. 私教还是枪手：基于通用大语言模型的计算机实践教学探索与反思[J]. 实验技术与管理，2024，41（5）：1-8.',
    '[2] 张金，宫晓利，高小鹏，等. 基于通用大语言模型的计算机系统创新实验设计[J]. 实验技术与管理，2024，41（10）：1-9.',
    '[3] 向尕，等. 新工科背景下“解决复杂工程问题”能力培养研究——以信息安全专业综合实习为例[J]. 软件导刊，2022，21（9）：211-218.',
    '[4] LangChain. LangGraph: building stateful multi-agent applications with LLMs[EB/OL]. (2024)[2026-07-31]. https://github.com/langchain-ai/langgraph.',
    '[5] WANG X, JI H, SHI C, et al. Heterogeneous graph attention network[C]//Proceedings of the World Wide Web Conference (WWW). 2019: 2022-2032.',
    '[6] SHINN N, CASSANO F, GOPINATH A, et al. Reflexion: language agents with verbal reinforcement learning[C]//Advances in Neural Information Processing Systems (NeurIPS). 2023.',
    '[7] WEBER M, DOMENICONI G, CHEN J, et al. Scalable graph learning for anti-money laundering: a first look[C]. 2018. DOI: 10.48550/arXiv.1812.00076.',
    '[8] WEBER M, DOMENICONI G, CHEN J, et al. Anti-money laundering in bitcoin: experimenting with graph convolutional networks for financial forensics[C]. 2019. DOI: 10.48550/arXiv.1908.02591.',
    '[9] HAMILTON W L, YING R, LESKOVEC J. Inductive representation learning on large graphs[C]//Advances in Neural Information Processing Systems (NeurIPS). 2017.',
    '[10] DOU Y, LIU Z, SUN L, et al. Enhancing graph neural networks by a label propagation algorithm for fraud detection[C]//Proceedings of the ACM International Conference on Information & Knowledge Management (CIKM). 2020: 2589-2592.',
    '[11] 蒋宗礼. 本科工程教育：聚焦学生解决复杂工程问题能力的培养[J]. 中国大学教学，2016（11）：27-30, 84.',
    '[12] 李剑. 网络空间安全专业研究生课程思政教育的探索与实践[J]. 信息安全研究，2024，10（2）：190-192.',
]
for r in refs:
    para(r, size=8.5, line=1.05, space_after=0)

doc.save(OUT)
print('已生成:', OUT)
